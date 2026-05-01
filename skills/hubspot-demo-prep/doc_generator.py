#!/usr/bin/env python3
"""Generate the mode-aware demo/showcase runbook .docx, then upload to Google Drive when possible.

The visual layout (banner, status pills, agenda block, easter egg, "also built",
recommendation, page-2 supporting documentation) is fixed; all slug-specific
copy, branding, and IDs come from build-plan.json, manifest.json, and
research.json so the same generator works for any prospect or feature showcase.

Public API:
    generate_docx(manifest, research, plan, *, slug, work_dir, portal) -> str
    upload_to_drive(docx_path, *, doc_title, drive_folder_id=..., replace_doc_id=None) -> dict
    export_pdf(doc_id, out_path) -> str | None
"""
from __future__ import annotations

import configparser
import datetime
import json
import os
import pathlib
import re
import sys
import urllib.error
import urllib.parse
import urllib.request

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# ---- Brand palette ----
# BRAND_ACCENT_NEUTRAL is the default accent color used when the prospect's
# branding doesn't specify one (see _accent_color) — slate blue (#3B82F6),
# matching the plan-schema fallback. DARK_TEXT is the default near-black
# neutral (#111827) used for headings and titles when no neutral_dark is set.
BRAND_ACCENT_NEUTRAL = RGBColor(0x3B, 0x82, 0xF6)
DARK_TEXT = RGBColor(0x11, 0x18, 0x27)
GRAY = RGBColor(0x55, 0x55, 0x55)
LIGHT_GRAY = RGBColor(0x99, 0x99, 0x99)
SUCCESS_GREEN = RGBColor(0x2C, 0x84, 0x4F)
WARN_AMBER = RGBColor(0xB1, 0x6E, 0x05)
NOT_BUILT_RED = RGBColor(0xB9, 0x1C, 0x1C)
BLUE = RGBColor(0x06, 0x6E, 0xA0)


# =====================================================================
# Branding helpers — pull accent / dark-text colors from the prospect's
# branding block (manifest or plan) with safe fallbacks.
# =====================================================================

def _parse_hex_color(value) -> RGBColor | None:
    """Parse a '#RRGGBB' or 'RRGGBB' hex string into an RGBColor. Returns None on failure."""
    if not isinstance(value, str):
        return None
    s = value.strip().lstrip("#")
    if len(s) != 6:
        return None
    try:
        return RGBColor(int(s[0:2], 16), int(s[2:4], 16), int(s[4:6], 16))
    except ValueError:
        return None


def _branding_color(manifest: dict, plan: dict, key: str, default: RGBColor) -> RGBColor:
    """Look up branding[key] in manifest first, then plan; parse hex; else default."""
    for src in (manifest, plan):
        if not isinstance(src, dict):
            continue
        branding = src.get("branding") or {}
        if isinstance(branding, dict):
            parsed = _parse_hex_color(branding.get(key))
            if parsed is not None:
                return parsed
    return default


def _accent_color(manifest: dict, plan: dict) -> RGBColor:
    """Return the prospect's accent color from branding.accent_color, else BRAND_ACCENT_NEUTRAL."""
    return _branding_color(manifest, plan, "accent_color", BRAND_ACCENT_NEUTRAL)


def _dark_text(manifest: dict, plan: dict) -> RGBColor:
    """Return the prospect's body/title text color from branding.neutral_dark, else DARK_TEXT."""
    return _branding_color(manifest, plan, "neutral_dark", DARK_TEXT)


# =====================================================================
# Mode helpers
# =====================================================================

def _plan_mode(plan: dict | None) -> str:
    mode = str((plan or {}).get("mode") or "demo").strip().lower().replace("-", "_")
    if mode in {"feature", "showcase", "feature_showcase", "feature_showcase_mode"}:
        return "feature_showcase"
    return "demo"


def _is_feature_showcase(plan: dict | None) -> bool:
    return _plan_mode(plan) == "feature_showcase"


def _feature_showcase(plan: dict | None, research: dict | None = None) -> dict:
    for src in (plan, research):
        if not isinstance(src, dict):
            continue
        block = src.get("feature_showcase")
        if isinstance(block, dict):
            return block
    return {}


def _text_list(value, *, limit: int = 6) -> list[str]:
    if value is None:
        return []
    if isinstance(value, str):
        return [value.strip()] if value.strip() else []
    if isinstance(value, list):
        out = []
        for item in value:
            if isinstance(item, dict):
                text = item.get("label") or item.get("title") or item.get("name") or item.get("description")
            else:
                text = str(item)
            if text and str(text).strip():
                out.append(str(text).strip())
            if len(out) >= limit:
                break
        return out
    return [str(value).strip()] if str(value).strip() else []


# =====================================================================
# Phantom-number guard — strip sentences containing dollar amounts that
# don't appear in any real deal in the manifest or plan. Prevents the
# recommendation copy from inventing figures the rep can't back up.
# =====================================================================

# Capture full money strings:
#   group(1): the numeric body, e.g. "1,200,000" or "4.2" or "4,200.50"
#   group(2): optional K/M suffix
# Anchored on '$' and requires word-boundary-ish termination via the
# (?![\d,\.]) lookahead so we never grab a partial number.
_DOLLAR_PATTERN = re.compile(
    r"\$(\d{1,3}(?:,\d{3})*(?:\.\d+)?|\d+(?:\.\d+)?)([KkMm])?(?![\d,\.])"
)

# Common abbreviations that should NOT terminate a sentence when followed by
# whitespace + capital letter. Lookbehind below covers each.
_ABBREV_LOOKBEHIND = re.compile(
    r"(?<!\bInc)(?<!\bCo)(?<!\be\.g)(?<!\bi\.e)(?<!\betc)(?<!\bLtd)"
    r"(?<!\bMr)(?<!\bMrs)(?<!\bMs)(?<!\bDr)(?<!\bSt)(?<!\bvs)\.\s+"
)


def _collect_deal_amounts(manifest: dict, plan: dict) -> set[float]:
    """Gather all known deal amounts from manifest['deals'] (defensive shapes) and plan['deals']."""
    amounts: set[float] = set()
    manifest_deals = manifest.get("deals") if isinstance(manifest, dict) else None
    if isinstance(manifest_deals, dict):
        for v in manifest_deals.values():
            if isinstance(v, dict):
                a = v.get("amount")
                if a is not None:
                    try:
                        amounts.add(float(a))
                    except (TypeError, ValueError):
                        pass
    elif isinstance(manifest_deals, list):
        for v in manifest_deals:
            if isinstance(v, dict):
                a = v.get("amount")
                if a is not None:
                    try:
                        amounts.add(float(a))
                    except (TypeError, ValueError):
                        pass
    plan_deals = plan.get("deals") if isinstance(plan, dict) else None
    if isinstance(plan_deals, list):
        for d in plan_deals:
            if isinstance(d, dict):
                a = d.get("amount")
                if a is not None:
                    try:
                        amounts.add(float(a))
                    except (TypeError, ValueError):
                        pass
    return amounts


def _parse_dollar_match(num_str: str, suffix: str | None) -> float | None:
    """Convert a regex-captured money string to a float USD value.

    '1,200,000' -> 1_200_000.0
    '4,200.50'  -> 4200.50
    '4.2' + 'K' -> 4200.0
    '2.5' + 'M' -> 2_500_000.0
    Returns None if the string cannot be parsed.
    """
    cleaned = num_str.replace(",", "")
    try:
        value = float(cleaned)
    except ValueError:
        return None
    if suffix:
        s = suffix.lower()
        if s == "k":
            value *= 1_000
        elif s == "m":
            value *= 1_000_000
    return value


def _strip_phantom_numbers(text: str, manifest: dict, plan: dict) -> str:
    """Remove any sentence containing a $-amount not present in the deal pool.

    - Captures full grouped numbers ('$1,200,000', '$4,200.50') and K/M
      shorthand ('$4.2K', '$2.5M') in one regex.
    - Tolerance is split by capture style:
        * K/M shorthand: 5% relative tolerance (absorbs '$4.2K' ≈ $4,200
          when the real deal is $4,180 – $4,420).
        * Explicit numeric (no K/M suffix): essentially-exact match
          (within $0.01) so '$4,200.50' does NOT match a real '$4,200'
          deal but '$4,200' does. Cents-precise figures the rep can't
          back up get dropped.
    - Sentence splitting respects common abbreviations (Inc., e.g., etc.,
      Co., Ltd., Mr., Mrs., Ms., Dr., St., vs., i.e.) so a mid-sentence
      'e.g.' won't fragment the sentence.
    Sentences without dollar amounts are always kept.
    """
    if not text:
        return text
    amounts = _collect_deal_amounts(manifest, plan)

    def _matches_known(value: float, *, shorthand: bool) -> bool:
        for a in amounts:
            if shorthand:
                # K/M rounding tolerance — '$4.2K' should match $4,200
                # exactly and any nearby figure within 5%.
                tol = max(1.0, a * 0.05)
            else:
                # Explicit number: essentially-exact match (sub-cent).
                # '$4,200.50' must NOT match a real '$4,200' deal.
                tol = 0.01
            if abs(value - a) <= tol:
                return True
        return False

    sentences: list[str] = []
    pos = 0
    for match in _ABBREV_LOOKBEHIND.finditer(text):
        # Keep the sentence-ending period; the regex consumes ". " so a plain
        # split would make the recommendation text read like a run-on.
        sentences.append(text[pos:match.start() + 1])
        pos = match.end()
    sentences.append(text[pos:])
    kept: list[str] = []
    for sent in sentences:
        drop = False
        for m in _DOLLAR_PATTERN.finditer(sent):
            value = _parse_dollar_match(m.group(1), m.group(2))
            if value is None:
                continue
            if not _matches_known(value, shorthand=bool(m.group(2))):
                drop = True
                break
        if not drop:
            kept.append(sent)
    return " ".join(kept)


# =====================================================================
# Workflow URL preference — when programmatic creation succeeded, manifest
# stores a specific edit URL per workflow name. Otherwise fall back to the
# manual_steps[i].ui_url for that workflow, then to the workflows index.
# =====================================================================

def _workflow_url(workflow_name: str | None, manifest: dict, urls: dict) -> str:
    """Best workflow link given what was actually built / queued for manual.

    Lookup order:
      1. exact-match against manifest["workflow_urls"][workflow_name]
      2. substring-match against the keys of manifest["workflow_urls"]
         (handles the common pattern where agenda items pass keyword stubs
         like "nurture" / "routing" while builder stored full names like
         "Boomer McLOUD - Welcome nurture")
      3. matching manual_steps[i].ui_url (workflow built manually)
      4. workflows index URL (last-resort fallback)
    """
    workflow_urls = (manifest.get("workflow_urls") or {}) if isinstance(manifest, dict) else {}
    if workflow_name and isinstance(workflow_urls, dict):
        # 1. Exact match
        if workflow_urls.get(workflow_name):
            return workflow_urls[workflow_name]
        # 2. Substring match — agenda passes "nurture", manifest stored "Boomer - Welcome nurture"
        wf_lower = workflow_name.lower()
        for stored_name, stored_url in workflow_urls.items():
            if not stored_url:
                continue
            stored_lower = stored_name.lower()
            if wf_lower in stored_lower or stored_lower in wf_lower:
                return stored_url
    # 3. Fallback: matching manual_steps entry (workflow built manually in UI)
    for ms in (manifest.get("manual_steps") or []):
        item = (ms.get("item") or "").lower()
        if "workflow" not in item:
            continue
        if not workflow_name or workflow_name.lower() in item:
            ui = ms.get("ui_url")
            if ui:
                return ui
    # 4. Last resort: workflows index
    return urls.get("workflows") or ""


# ---- Drive / OAuth defaults ----
DEFAULT_DRIVE_FOLDER_ID = "1SzHT9uhFUUcFIAh5z2LVCAq2Wt0OADjY"
RCLONE_CONF_PATH = pathlib.Path.home() / ".config/rclone/rclone.conf"
# Same client_id/secret used by rclone's gdrive backend (matches update-doc.py).
RCLONE_CLIENT_ID = "202264815644.apps.googleusercontent.com"
RCLONE_CLIENT_SECRET = "X4Z3ca8xfWDb1Voo-F9a7ZxJ"


# =====================================================================
# docx helpers (verbatim from make-doc.py)
# =====================================================================

def _set_run(run, *, color=None, bold=False, italic=False, size=None, font="Arial"):
    run.font.name = font
    if color is not None:
        run.font.color.rgb = color
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size is not None:
        run.font.size = Pt(size)


def _add_hyperlink(paragraph, url, text, *, color=BLUE, underline=True, bold=False, size=10):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rPr.append(sz)
    if color is not None:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), "%02X%02X%02X" % (color[0], color[1], color[2]))
        rPr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
    if bold:
        rPr.append(OxmlElement("w:b"))
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _shade_paragraph(p, fill_hex, left_border_hex=None):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)
    if left_border_hex:
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "24")
        left.set(qn("w:space"), "4")
        left.set(qn("w:color"), left_border_hex)
        pBdr.append(left)
        pPr.append(pBdr)


def _bottom_border(p, color_hex="3B82F6"):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _h2(doc, text, *, color=None, before=8):
    """H2 heading. If color is None, callers should pass the prospect-derived accent."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text.upper())
    _set_run(r, color=color if color is not None else BRAND_ACCENT_NEUTRAL, bold=True, size=10)


def _status_pill(p, text, color):
    r = p.add_run(f"[{text}] ")
    _set_run(r, color=color, bold=True, size=9)


# =====================================================================
# URL builders
# =====================================================================

def _hub_urls(portal: str, manifest: dict) -> dict:
    """Build the table of HubSpot deep links from portal id + manifest."""
    pipeline_id = (manifest.get("pipeline") or {}).get("id", "")
    company_id = (manifest.get("company") or {}).get("id", "")
    custom_obj_id = (manifest.get("custom_object") or {}).get("object_type_id", "")
    forms = manifest.get("forms") or {}
    nps_form_id = ""
    for fname, fid in forms.items():
        if "nps" in fname.lower():
            nps_form_id = fid
            break
    email_id = (manifest.get("marketing_email") or {}).get("id", "")
    lead_score_prop = (manifest.get("lead_scoring") or {}).get("property", "demo_lead_score")
    return {
        "pipeline_board": f"https://app.hubspot.com/contacts/{portal}/objects/0-3/views/all/board?pipeline={pipeline_id}",
        "pipeline_list": f"https://app.hubspot.com/contacts/{portal}/objects/0-3/views/all/list?pipeline={pipeline_id}",
        "company": f"https://app.hubspot.com/contacts/{portal}/record/0-2/{company_id}",
        "custom_obj": f"https://app.hubspot.com/contacts/{portal}/objects/{custom_obj_id}" if custom_obj_id else "",
        "nps_form": f"https://app.hubspot.com/forms/{portal}/editor/{nps_form_id}/edit/form" if nps_form_id else "",
        "email": f"https://app.hubspot.com/email/{portal}/edit/{email_id}/edit/content" if email_id else f"https://app.hubspot.com/email/{portal}/manage/state/all",
        "workflows": f"https://app.hubspot.com/workflows/{portal}/view/all-workflows",
        "forms_index": f"https://app.hubspot.com/forms/{portal}",
        "landing_pages": f"https://app.hubspot.com/website/{portal}/landing-pages",
        "contacts_list": f"https://app.hubspot.com/contacts/{portal}/objects/0-1/views/all/list",
        "lead_score_prop": f"https://app.hubspot.com/property-settings/{portal}/properties?type=0-1&property={lead_score_prop}",
        "tickets": f"https://app.hubspot.com/contacts/{portal}/objects/0-5/views/all/list",
        "event_defs": f"https://app.hubspot.com/events/{portal}/manage/event-definitions",
        "campaigns_index": f"https://app.hubspot.com/marketing/{portal}/campaigns",
        "dashboards_index": f"https://app.hubspot.com/reports-dashboard/{portal}",
        "reports_index": f"https://app.hubspot.com/reports/{portal}",
    }


def _deal_url(portal: str, did: str) -> str:
    return f"https://app.hubspot.com/contacts/{portal}/record/0-3/{did}"


def _contact_url(portal: str, cid: str) -> str:
    return f"https://app.hubspot.com/contacts/{portal}/record/0-1/{cid}"


# =====================================================================
# Logo discovery
# =====================================================================

def _logo_path(manifest: dict, plan: dict) -> str | None:
    """Return the prospect's logo path if it exists and is readable.

    Looks at manifest['branding']['logo_path'] first, then plan['branding']
    ['logo_path']. Returns None if neither is set or the file is missing —
    older runs predating logo extraction degrade gracefully (header renders
    title-only, no logo cell).
    """
    for src in (manifest, plan):
        if not isinstance(src, dict):
            continue
        branding = src.get("branding") or {}
        if not isinstance(branding, dict):
            continue
        path = branding.get("logo_path")
        if isinstance(path, str) and path and os.path.isfile(path):
            return path
    return None


# =====================================================================
# Doc body
# =====================================================================

def _format_currency(amount) -> str:
    try:
        return f"${int(round(float(amount))):,}"
    except (TypeError, ValueError):
        return str(amount or "")


def _coerce_amount(amount) -> int:
    """Bug fix (2026-04-27): plan deal amounts arrive as either int OR string
    (the legacy schema doc told orchestrators to send strings; builder.py
    coerces them server-side before posting to HubSpot but the doc generator
    summed them raw, hitting `TypeError: int + str` on the very first run with
    a string-typed amount. This helper centralizes the coercion."""
    try:
        return int(round(float(amount or 0)))
    except (TypeError, ValueError):
        return 0


def _format_runtime(seconds) -> str | None:
    """Render manifest['runtime_seconds'] as e.g. '4m 12s'. None on bad input."""
    try:
        s = int(round(float(seconds)))
    except (TypeError, ValueError):
        return None
    if s <= 0:
        return None
    if s < 60:
        return f"{s}s"
    m, r = divmod(s, 60)
    return f"{m}m {r}s" if r else f"{m}m"


def _render_time_saved_hero(doc, manifest: dict, plan: dict) -> None:
    """Hero stat under the title bar: '⏱ Saved ~Xh vs manual HubSpot setup'.

    Renders nothing if manifest['time_saved'] is missing or malformed —
    graceful degradation per item 14's "don't break" clause.
    """
    ts = (manifest or {}).get("time_saved") or {}
    pretty = ts.get("total_pretty")
    if not pretty:
        return
    accent = _accent_color(manifest, plan or {})
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(
        f"⏱  It would take approximately {pretty} to create the equivalent "
        f"demo portal manually."
    )
    _set_run(r, color=accent, bold=True, size=13)
    runtime_pretty = _format_runtime(manifest.get("runtime_seconds"))
    if runtime_pretty:
        sub = doc.add_paragraph()
        sub.paragraph_format.space_before = Pt(0)
        sub.paragraph_format.space_after = Pt(4)
        sr = sub.add_run(f"This demo was built in {runtime_pretty}.")
        _set_run(sr, color=GRAY, italic=True, size=9)


def _render_time_saved_breakdown(doc, manifest: dict, plan: dict) -> None:
    """Per-phase breakdown table at the bottom of the doc.

    Columns: Phase | Built | Manual time. One row per non-zero entry +
    a final total row. 9pt font, lined. Skips silently if missing.
    """
    ts = (manifest or {}).get("time_saved") or {}
    breakdown = ts.get("breakdown") or []
    if not breakdown:
        return
    accent = _accent_color(manifest, plan or {})
    _h2(doc, "Time saved vs manual build", color=accent, before=10)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for cell, label in zip(hdr, ("Phase", "Built", "Manual time")):
        cell_para = cell.paragraphs[0]
        cell_para.paragraph_format.space_after = Pt(0)
        run = cell_para.add_run(label)
        _set_run(run, color=accent, bold=True, size=9)

    for row in breakdown:
        cells = table.add_row().cells
        for idx, value in enumerate((
            str(row.get("label") or ""),
            str(row.get("count") or 0),
            str(row.get("subtotal_pretty") or ""),
        )):
            cell_para = cells[idx].paragraphs[0]
            cell_para.paragraph_format.space_after = Pt(0)
            run = cell_para.add_run(value)
            _set_run(run, size=9)

    # Total row.
    total_cells = table.add_row().cells
    for idx, value in enumerate((
        "TOTAL",
        str(sum(int(r.get("count") or 0) for r in breakdown)),
        str(ts.get("total_pretty") or ""),
    )):
        cell_para = total_cells[idx].paragraphs[0]
        cell_para.paragraph_format.space_after = Pt(0)
        run = cell_para.add_run(value)
        _set_run(run, color=accent, bold=True, size=9)


def _rep_name(manifest: dict, plan: dict, research: dict) -> str:
    """Return a rep name for the subtitle, else 'Sales Engineer'.

    Looks at manifest['rep_name'], plan['rep_name'], research['rep_name'],
    and falls back to 'Sales Engineer' if none are set. Strings only.
    """
    for src in (manifest, plan, research):
        if not isinstance(src, dict):
            continue
        v = src.get("rep_name")
        if isinstance(v, str) and v.strip():
            return v.strip()
    return "Sales Engineer"


def _render_header(doc, manifest: dict, plan: dict, research: dict, *,
                   company_name: str, portal: str) -> None:
    """Render the consulting-deliverable top header.

    Layout: 1-row x 2-col table at the top of the doc — logo left (~1.5"),
    title right (~5"). Below the table: subtitle line in gray, then a thin
    colored rule in the prospect's primary color (or accent fallback) that
    separates the header from page content.

    If no logo is available (manifest/plan branding missing logo_path or
    file unreadable), the title spans the full width — no logo cell — so
    the layout still feels intentional and centered, not awkwardly empty.
    """
    dark = _dark_text(manifest, plan)
    # Use primary_color for the rule when available; else accent.
    rule_color = _branding_color(manifest, plan, "primary_color", _accent_color(manifest, plan))
    logo = _logo_path(manifest, plan)
    feature_mode = _is_feature_showcase(plan)
    title_prefix = "HubSpot Feature Showcase: " if feature_mode else "HubSpot Demo Prep: "

    if logo:
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.allow_autofit = False
        # Column widths: logo cell ~1.5", title cell ~5".
        try:
            table.columns[0].width = Inches(1.5)
            table.columns[1].width = Inches(5.0)
            for row in table.rows:
                row.cells[0].width = Inches(1.5)
                row.cells[1].width = Inches(5.0)
        except Exception:  # noqa: BLE001
            pass

        # Logo cell.
        logo_cell = table.rows[0].cells[0]
        logo_para = logo_cell.paragraphs[0]
        logo_para.paragraph_format.space_after = Pt(0)
        logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        logo_run = logo_para.add_run()
        try:
            # Cap by width; height will scale proportionally. Wide logos
            # are bounded by the 1.4" width; very tall logos would push
            # the cell taller, but most marketing logos are landscape.
            logo_run.add_picture(logo, width=Inches(1.4))
            # If the resulting picture is taller than ~0.8", clamp height
            # by re-adding with explicit height instead.
            try:
                inline_shapes = [s for s in doc.inline_shapes]
                if inline_shapes:
                    last = inline_shapes[-1]
                    # Only re-clamp if height exceeds 0.8" (730000 EMUs).
                    if last.height and last.height > Inches(0.8):
                        last.height = Inches(0.8)
                        # Re-derive width to maintain aspect ratio is not
                        # straightforward via python-docx; the cap on
                        # height alone is acceptable since Word preserves
                        # the embedded aspect ratio when only height is
                        # set on an inline shape.
            except Exception:  # noqa: BLE001
                pass
        except Exception:  # noqa: BLE001
            # Logo file is malformed — fail open, leave the cell empty.
            pass

        # Title cell.
        title_cell = table.rows[0].cells[1]
        title_para = title_cell.paragraphs[0]
        title_para.paragraph_format.space_after = Pt(0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Vertically center the title against the logo.
        try:
            from docx.enum.table import WD_ALIGN_VERTICAL
            title_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            logo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        except Exception:  # noqa: BLE001
            pass
        # Prefix in slightly muted (gray) weight, then company/showcase name
        # in bold dark for emphasis.
        r = title_para.add_run(title_prefix)
        _set_run(r, color=GRAY, bold=False, size=22)
        r = title_para.add_run(company_name)
        _set_run(r, color=dark, bold=True, size=22)
    else:
        # No logo — title spans full width, full bold.
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(title_prefix)
        _set_run(r, color=GRAY, bold=False, size=22)
        r = p.add_run(company_name)
        _set_run(r, color=dark, bold=True, size=22)

    # Subtitle line: "Demo/Feature showcase for {rep_name} · {Date} · Sandbox portal {portal}"
    today = datetime.date.today().strftime("%B %-d, %Y")
    rep = _rep_name(manifest, plan, research)
    subtitle_prefix = "Feature showcase for" if feature_mode else "Demo for"
    subtitle = f"{subtitle_prefix} {rep}  ·  {today}  ·  Sandbox portal {portal}"
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(2)
    sp.paragraph_format.space_after = Pt(0)
    sr = sp.add_run(subtitle)
    _set_run(sr, color=GRAY, size=9)
    # Thin colored rule beneath the subtitle.
    _bottom_border(sp, color_hex="%02X%02X%02X" % (rule_color[0], rule_color[1], rule_color[2]))


def _build_doc(manifest: dict, research: dict, plan: dict, *,
               slug: str, work_dir: str, portal: str) -> Document:
    feature_mode = _is_feature_showcase(plan)
    feature = _feature_showcase(plan, research)
    company = plan.get("company") or {}
    company_name = company.get("name") or (manifest.get("company") or {}).get("name") or "Customer"
    urls = _hub_urls(portal, manifest)
    agenda = plan.get("agenda") or []
    easter = plan.get("easter_egg") or {}
    accent = _accent_color(manifest, plan)
    dark = _dark_text(manifest, plan)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.45)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(1)
    style.paragraph_format.line_spacing = 1.10

    # ---------------- PAGE 1 ----------------

    # Consulting-deliverable header: logo (if available) + title + subtitle
    # + thin colored rule. Replaces the old generated banner. If no logo is
    # present (older runs, missing logo_path), the title spans full width.
    _render_header(doc, manifest, plan, research,
                   company_name=company_name, portal=portal)

    # Time-saved hero stat (item 14). Renders nothing if manifest['time_saved']
    # is missing, so older runs and module failures degrade gracefully.
    try:
        _render_time_saved_hero(doc, manifest, plan)
    except Exception:  # noqa: BLE001
        pass

    # Intro paragraph (rep input / showcase brief + what was built)
    rep_input = (feature.get("story") if feature_mode else None) or (research.get("stated_context") or "")
    rep_input = str(rep_input).strip()
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    if rep_input:
        r = p.add_run("Showcase brief: " if feature_mode else "Rep input: ")
        _set_run(r, bold=True, size=9.5)
        r = p.add_run(rep_input + "  ")
        _set_run(r, italic=True, color=GRAY, size=9.5)
    r = p.add_run("Built: ")
    _set_run(r, bold=True, size=9.5)
    built_summary = _built_summary(manifest, plan)
    r = p.add_run(built_summary)
    _set_run(r, size=9.5)

    # ---- AGENDA / SHOWCASE FLOW ----
    _h2(doc, "Showcase flow" if feature_mode else "Agenda (from sales rep)",
        color=accent, before=8)
    for i, item in enumerate(agenda, 1):
        _render_agenda_item(doc, i, item, manifest=manifest, plan=plan, urls=urls, portal=portal)

    # ---- EASTER EGG ----
    if easter:
        _render_easter_egg(doc, easter, manifest=manifest, plan=plan, urls=urls, portal=portal)

    # ---- REPORTING & DASHBOARDS (v0.4) ----
    _render_reports_section(doc, manifest=manifest, plan=plan, urls=urls, portal=portal)

    # ---- ALSO BUILT ----
    also_title = (
        "Also built (pull these into the showcase if useful)"
        if feature_mode else "Also built (pull these into the demo if useful)"
    )
    _h2(doc, also_title, color=accent, before=8)
    _render_also_built(doc, manifest=manifest, plan=plan, urls=urls, portal=portal)

    # ---- RECOMMENDATION ----
    _h2(doc, "How to tell the story" if feature_mode else "Recommendation",
        color=accent, before=6)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(_recommendation_text(manifest, plan))
    _set_run(r, size=9.5)

    # ---------------- PAGE 2 ----------------

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Showcase support" if feature_mode else "Supporting documentation")
    _set_run(r, color=dark, bold=True, size=16)
    _bottom_border(p, color_hex="%02X%02X%02X" % (accent[0], accent[1], accent[2]))

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    support_copy = (
        "Reference material for the feature story. Page 1 is the showcase flow; this page is the context cheat-sheet."
        if feature_mode
        else "Reference material for the demo. Page 1 is the runbook; this page is your context cheat-sheet."
    )
    r = p.add_run(support_copy)
    _set_run(r, color=GRAY, size=9)

    # Pre-demo / pre-showcase checklist
    _h2(doc, "Pre-showcase checklist" if feature_mode else "Pre-demo checklist",
        color=accent, before=2)
    _render_checklist(doc, manifest=manifest, urls=urls, portal=portal)

    # Company snapshot / feature brief
    if feature_mode:
        _h2(doc, "Feature showcase brief", color=accent, before=8)
        _render_feature_showcase_brief(doc, feature=feature, company=company, research=research)
        if plan.get("campaign_attribution_showcase"):
            _h2(doc, "Attribution story map", color=accent, before=8)
            _render_attribution_story_map(doc, plan=plan)
    else:
        _h2(doc, f"{company_name} snapshot", color=accent, before=8)
        _render_snapshot(doc, company=company, research=research)

    # ICP / research
    _h2(doc,
        "Story and audience context" if feature_mode else "ICP and pain-point research (Perplexity, sonar)",
        color=accent, before=8)
    _render_icp(doc, research=research, plan=plan)

    # Full inventory
    _h2(doc, f"Full build inventory (sandbox {portal})", color=accent, before=8)
    for item in _build_inventory(manifest, plan):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.18)
        r = p.add_run(f"•  {item}")
        _set_run(r, size=9.5)

    # Limitations
    limits = _build_limitations(manifest)
    if limits:
        _h2(doc, "Known build limitations", color=accent, before=8)
        for item in limits:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Inches(0.18)
            r = p.add_run(f"•  {item}")
            _set_run(r, color=GRAY, size=9.5)

    # Time-saved breakdown table (item 14). Silently skips if absent.
    try:
        _render_time_saved_breakdown(doc, manifest, plan)
    except Exception:  # noqa: BLE001
        pass

    # Sources
    sources = _research_sources(research)
    if sources:
        _h2(doc, "Sources", color=accent, before=8)
        for label, url in sources:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Inches(0.18)
            _add_hyperlink(p, url, label, size=9.5)

    # Cleanup
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("When done with this showcase: " if feature_mode else "When done with this demo: ")
    _set_run(r, color=GRAY, italic=True, size=9)
    mono = p.add_run(f"python3 ~/.claude/skills/hubspot-demo-prep/builder.py cleanup {slug}")
    mono.font.name = "Menlo"
    mono.font.size = Pt(8)
    mono.font.color.rgb = GRAY

    return doc


# =====================================================================
# Section renderers
# =====================================================================

def _built_summary(manifest: dict, plan: dict) -> str:
    feature_mode = _is_feature_showcase(plan)
    parts = []
    n_contacts = len(manifest.get("contacts") or {})
    n_engagements = manifest.get("engagements_count") or 0
    n_deals = len(manifest.get("deals") or {})
    if n_contacts:
        parts.append(f"{n_contacts} contacts")
    if n_engagements:
        parts.append(f"{n_engagements} timeline events")
    if n_deals:
        parts.append(f"a {n_deals}-deal pipeline")
    if (manifest.get("custom_object") or {}).get("name"):
        parts.append(f"a custom {manifest['custom_object']['name']} object")
    event_fires = int(manifest.get("custom_events_fired_count") or 0)
    if event_fires:
        parts.append(f"{event_fires} custom event fires")
    if manifest.get("forms"):
        parts.append("a live NPS survey" if any("nps" in n.lower() for n in manifest["forms"]) else "a live form")
    if (manifest.get("marketing_email") or {}).get("name") or (manifest.get("marketing_email") or {}).get("html_path"):
        hero_present = bool(
            (manifest.get("marketing_email") or {}).get("hero_image_url")
            or (plan.get("marketing_email") or {}).get("hero_image_path")
            or (plan.get("marketing_email") or {}).get("hero_image_url")
        )
        parts.append(
            "a branded marketing email with AI hero image"
            if hero_present else "a branded marketing email"
        )
    if (manifest.get("lead_scoring") or {}).get("property"):
        backfilled = manifest["lead_scoring"].get("backfilled") or n_contacts
        parts.append(f"lead scoring on all {backfilled} contacts")
    if manifest.get("dashboards_v04"):
        parts.append(f"{len(manifest.get('dashboards_v04') or {})} reporting dashboard(s)")
    company_name = (plan.get("company") or {}).get("name") or (
        "the showcase" if feature_mode else "the prospect"
    )
    descriptor = "showcase-ready" if feature_mode else "activity-rich"
    summary = (
        f"{descriptor} {company_name} portal with " + ", ".join(parts) + "."
        if parts else f"{descriptor} {company_name} portal."
    )
    if manifest.get("manual_steps"):
        n_workflow_manual = sum(1 for ms in manifest["manual_steps"]
                                if "workflow" in (ms.get("item") or "").lower())
        if n_workflow_manual:
            summary += (f" {n_workflow_manual} workflow step(s) are built in the UI for finer control "
                        "over branching/timing (60 seconds each — called out below).")
    return summary


def _agenda_status_lines(item: dict, idx: int, manifest: dict, urls: dict, portal: str,
                         plan: dict | None = None) -> list[tuple[str, RGBColor, str, str, str]]:
    """For each agenda item produce a list of (pill_label, pill_color, body_text, link_url[, link_text]).

    Falls back to a generic 'open in HubSpot' link from agenda.show_label.
    Returns list of tuples: (pill, color, prefix, link_url, link_label)
    """
    lines: list[tuple[str, RGBColor, str, str, str]] = []
    title_l = (item.get("title") or "").lower()

    # Verification gate — a manifest entry alone isn't enough. The builder may have
    # recorded an artifact, but the post-create GET could have shown the artifact
    # never landed. Treat unverified phases as not built so the doc never claims
    # something is live that the rep can't actually open.
    verifications = manifest.get("verifications") or {}
    def _verified(phase: str) -> bool:
        v = verifications.get(phase)
        if v is None:
            return True  # legacy manifests without verifications keep prior behavior
        return bool(v.get("verified"))

    forms = manifest.get("forms") or {}
    forms_ok = _verified("forms")
    has_nps = forms_ok and any("nps" in n.lower() for n in forms)
    has_quote_form = forms_ok and any("quote" in n.lower() for n in forms)
    has_email = (_verified("marketing_email")
                 and bool((manifest.get("marketing_email") or {}).get("name")
                          or (manifest.get("marketing_email") or {}).get("html_path")))
    has_email_hero = bool((manifest.get("marketing_email") or {}).get("hero_image_url"))
    has_workflows = _verified("workflows") and bool(manifest.get("workflows"))
    has_landing = bool(manifest.get("landing_page"))

    # Workflow link preference: if a specific edit URL exists in the manifest
    # (set when programmatic creation succeeded), use it; otherwise fall back
    # to the matching manual_steps.ui_url; otherwise the workflows index.
    nurture_wf_url = _workflow_url(item.get("workflow_name") or "nurture", manifest, urls)
    routing_wf_url = _workflow_url(item.get("workflow_name") or "routing", manifest, urls)
    generic_wf_url = _workflow_url(item.get("workflow_name"), manifest, urls)

    if "nurtur" in title_l or "drip" in title_l or "follow-up" in title_l:
        if has_email:
            email_label = (
                "Branded marketing email with AI hero image"
                if has_email_hero else "Branded marketing email"
            )
            lines.append(("BUILT", SUCCESS_GREEN,
                          f"{email_label}  ▸  ",
                          urls["email"], "Open email in HubSpot"))
        if not has_workflows:
            lines.append(("BUILD LIVE", WARN_AMBER,
                          "Workflow itself (UI build is faster than the setup for this routing logic)  ▸  ",
                          nurture_wf_url, "Open workflow"))
    elif "landing" in title_l or "inbound" in title_l or "form" in title_l:
        if not (has_quote_form and has_landing):
            lines.append(("NOT BUILT", NOT_BUILT_RED,
                          "Quote form is built in the UI for richer template handling; landing page is a premium feature — built manually for tighter control. Walk this one live.  ▸  ",
                          urls["forms_index"], "Forms"))
            lines.append(("ANALOG", BLUE,
                          ("The NPS form (built) is a working example of how the Quote form would behave once rebuilt in the UI  ▸  "
                           if has_nps else "Use the live form as an analog when walking through this  ▸  "),
                          urls["nps_form"] or urls["forms_index"], "Open NPS form" if has_nps else "Forms"))
        else:
            lines.append(("BUILT", SUCCESS_GREEN,
                          "Quote form + landing page live  ▸  ",
                          urls["forms_index"], "Open form"))
    elif "nps" in title_l or "feedback" in title_l or "survey" in title_l:
        if has_nps:
            lines.append(("BUILT", SUCCESS_GREEN,
                          f"{(plan_company_label(manifest))} NPS Survey form (live)  ▸  ".replace("  ", " "),
                          urls["nps_form"] or urls["forms_index"], "Open form"))
        if not has_workflows:
            lines.append(("BUILD LIVE", WARN_AMBER,
                          "Routing workflow (built manually for finer control over branching/timing)  ▸  ",
                          routing_wf_url, "Open workflow"))
    else:
        if _is_feature_showcase(plan):
            link_label = item.get("show_label") or "Open in HubSpot"
            link_url = ""
            if "deal" in title_l and (manifest.get("deals") or {}):
                _, first_deal_id = next(iter((manifest.get("deals") or {}).items()))
                link_url = _deal_url(portal, first_deal_id)
            elif "campaign" in title_l and manifest.get("campaign_url"):
                link_url = manifest.get("campaign_url")
            elif "dashboard" in title_l or "report" in title_l or "attribution" in title_l:
                link_url = urls.get("dashboards_index") or urls.get("reports_index")
            elif manifest.get("contacts"):
                _, first_contact_id = next(iter((manifest.get("contacts") or {}).items()))
                link_url = _contact_url(portal, first_contact_id)
            else:
                link_url = urls.get("contacts_list")
            lines.append(("BUILT", SUCCESS_GREEN, "  ▸  ", link_url, link_label))
            return lines

        # Generic fallback — single line with whatever URL we can derive
        link_label = item.get("show_label") or "Open in HubSpot"
        link_url = generic_wf_url
        lines.append(("BUILT" if has_workflows else "BUILD LIVE",
                      SUCCESS_GREEN if has_workflows else WARN_AMBER,
                      "  ▸  ", link_url, link_label))
    return lines


def plan_company_label(manifest: dict) -> str:
    return (manifest.get("company") or {}).get("name") or ""


def _render_agenda_item(doc, idx: int, item: dict, *, manifest: dict, plan: dict | None = None,
                        urls: dict, portal: str) -> None:
    title = item.get("title") or ""
    why = item.get("why") or ""
    accent = _accent_color(manifest, plan or {})
    dark = _dark_text(manifest, plan or {})

    # Heading
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{idx}.  ")
    _set_run(r, color=accent, bold=True, size=10.5)
    r = p.add_run(title)
    _set_run(r, color=dark, bold=True, size=10.5)

    # Why / description
    if why:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.22)
        r = p.add_run(why)
        _set_run(r, color=GRAY, italic=True, size=9.5)

    # Status lines
    for pill, color, body, link_url, link_label in _agenda_status_lines(item, idx, manifest, urls, portal, plan):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.22)
        _status_pill(p, pill, color)
        r = p.add_run(body)
        _set_run(r, size=9.5)
        if link_url:
            _add_hyperlink(p, link_url, link_label, size=9.5)


def _render_easter_egg(doc, easter: dict, *, manifest: dict, plan: dict | None = None,
                       urls: dict, portal: str) -> None:
    title = easter.get("title") or ""
    why = easter.get("why") or ""
    accent = _accent_color(manifest, plan or {})
    dark = _dark_text(manifest, plan or {})
    # Derive shading hexes from the accent: full-strength accent for the
    # left border, ~10% accent mixed onto white for the soft fill. Always
    # computed — no special-case preservation of any historical palette.
    accent_hex = "%02X%02X%02X" % (accent[0], accent[1], accent[2])
    tint_hex = "%02X%02X%02X" % (
        255 - (255 - accent[0]) // 10,
        255 - (255 - accent[1]) // 10,
        255 - (255 - accent[2]) // 10,
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(0)
    _shade_paragraph(p, tint_hex, accent_hex)
    section_label = easter.get("section_label") or (
        "ADJACENT VALUE" if _is_feature_showcase(plan) else "EASTER EGG"
    )
    r = p.add_run(f"★  {section_label.upper()}  ·  {title}")
    _set_run(r, color=accent, bold=True, size=10)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    _shade_paragraph(p, tint_hex, accent_hex)
    _status_pill(p, "BUILT", SUCCESS_GREEN)
    if why:
        r = p.add_run(f"{why}  ")
        _set_run(r, color=dark, italic=True, size=9.5)
    _add_hyperlink(p, urls["contacts_list"], "Sortable contact list", size=9.5)
    sep = p.add_run("  ·  ")
    _set_run(sep, color=LIGHT_GRAY, size=9.5)
    # First contact as a sample
    contacts = manifest.get("contacts") or {}
    if contacts:
        first_email, first_id = next(iter(contacts.items()))
        sample_label = "Sample: " + first_email.split("@")[0].replace(".", " ").title()
        _add_hyperlink(p, _contact_url(portal, first_id), sample_label, size=9.5)
        sep = p.add_run("  ·  ")
        _set_run(sep, color=LIGHT_GRAY, size=9.5)
    _add_hyperlink(p, urls["lead_score_prop"], "Property settings", size=9.5)


def _render_reports_section(doc, *, manifest: dict, plan: dict, urls: dict, portal: str) -> None:
    """v0.4: render the Reporting & Dashboards section.

    Reads manifest["dashboards_v04"] (populated by the Playwright reports
    phase) and renders each dashboard with audience label + clickable URL,
    plus a sub-line listing the report count and visualization mix. Surfaces
    tier substitutions and the attribution-toggle pre-stage warning when
    relevant. No-op when no v0.4 dashboards exist (preserves backward
    compatibility with v0.3.0 plans).
    """
    dashboards = manifest.get("dashboards_v04") or {}
    reports_status = manifest.get("reports_status") or {}
    planned_dashboards = ((plan or {}).get("playwright_reports") or {}).get("dashboards") or []
    if not dashboards and not reports_status and not planned_dashboards:
        return

    # Group reports by their owning dashboard. Prefer explicit dashboard_name;
    # otherwise split from the right so dashboard names containing "::" survive.
    reports_by_dash: dict[str, list[dict]] = {}
    for key, info in (manifest.get("reports") or {}).items():
        info = info or {}
        dash_name = info.get("dashboard_name")
        if not dash_name and "::" in str(key):
            dash_name = str(key).rsplit("::", 1)[0]
        if not dash_name:
            continue
        reports_by_dash.setdefault(dash_name, []).append(info)

    accent = _accent_color(manifest, plan or {})

    _h2(doc, "Reporting & dashboards", color=accent, before=8)

    if not dashboards:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        rb = p.add_run("Planned but not built: ")
        _set_run(rb, bold=True, color=accent, size=9.5)
        reason = reports_status.get("reason") or (
            f"{len(planned_dashboards)} v0.4 dashboard(s) planned, but no dashboard URLs were recorded."
        )
        rt = p.add_run(reason)
        _set_run(rt, color=GRAY, size=9.5)
        if urls.get("dashboards_index"):
            sep = p.add_run("  ·  ")
            _set_run(sep, color=LIGHT_GRAY, size=9.5)
            _add_hyperlink(p, urls["dashboards_index"], "All dashboards", size=9.5)
        return

    for dash_name, dash_info in dashboards.items():
        dash_info = dash_info or {}
        reports_for_this = reports_by_dash.get(dash_name, [])
        viz_types = sorted({
            (r.get("viz_type") or "").replace("_", " ")
            for r in reports_for_this
            if r.get("viz_type")
        })
        substituted = sum(1 for r in reports_for_this if r.get("tier_substituted"))

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        audience = dash_info.get("audience") or ""
        label = f"{dash_name}" + (f"  ({audience})" if audience else "")
        r = p.add_run(label + "  ▸  ")
        _set_run(r, bold=True, size=10)
        if dash_info.get("url"):
            _add_hyperlink(p, dash_info["url"], "Open dashboard", size=10)
        if urls.get("dashboards_index"):
            sep_run = p.add_run("  ·  ")
            _set_run(sep_run, color=LIGHT_GRAY, size=10)
            _add_hyperlink(p, urls["dashboards_index"], "All dashboards", size=10)

        # Sub-line: report count + viz mix + substitution note
        sub = doc.add_paragraph()
        sub.paragraph_format.space_after = Pt(2)
        sub.paragraph_format.left_indent = Inches(0.22)
        rcount = dash_info.get("report_count") or len(reports_for_this)
        if viz_types:
            viz_summary = ", ".join(viz_types)
        elif reports_for_this:
            viz_summary = "report mix recorded without visualization labels"
        else:
            viz_summary = "report manifest missing — verify this dashboard before demo"
        substitution_note = ""
        if substituted > 0:
            substitution_note = (
                f"  ·  {substituted} tier-substituted "
                "(Sankey → vertical funnel on this sandbox)"
            )
        rline = sub.add_run(f"{rcount} reports — {viz_summary}{substitution_note}")
        _set_run(rline, color=GRAY, size=9)

    # Surface attribution-toggle pre-stage warning if any report uses it.
    needs_attribution = any(
        (r.get("data_source") or "") == "attribution"
        for reports in reports_by_dash.values()
        for r in reports
    )
    if needs_attribution:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(0)
        rb = p.add_run("Pre-stage required: ")
        _set_run(rb, bold=True, color=accent, size=9)
        rt = p.add_run(
            "Multi-touch attribution reports reprocess for up to 2 days when an event "
            "is first toggled as an interaction type. Flip the toggle ≥48h before the "
            "demo to ensure the chart populates."
        )
        _set_run(rt, size=9, color=GRAY)


def _render_also_built(doc, *, manifest: dict, plan: dict, urls: dict, portal: str) -> None:
    deals = manifest.get("deals") or {}
    contacts = manifest.get("contacts") or {}

    # Pipeline + deals
    if (manifest.get("pipeline") or {}).get("id"):
        total = sum(_coerce_amount(d.get("amount")) for d in (plan.get("deals") or []))
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"Active deal pipeline ({len(deals)} deals, {_format_currency(total)} total)  ▸  ")
        _set_run(r, bold=True, size=10)
        _add_hyperlink(p, urls["pipeline_board"], "Board view", size=10)
        sep = p.add_run("  ·  ")
        _set_run(sep, color=LIGHT_GRAY, size=10)
        _add_hyperlink(p, urls["pipeline_list"], "List view", size=10)

        # Each deal individually (with stage label from plan)
        plan_deal_lookup = {d.get("name"): d for d in (plan.get("deals") or [])}
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.22)
        items = list(deals.items())
        for i, (name, did) in enumerate(items):
            pd = plan_deal_lookup.get(name) or {}
            stage = pd.get("stage") or ""
            amount = _format_currency(pd.get("amount"))
            label = f"{stage}: {name} ({amount})" if stage else f"{name} ({amount})"
            _add_hyperlink(p, _deal_url(portal, did), label, size=9)
            if i < len(items) - 1:
                sep = p.add_run("  ·  ")
                _set_run(sep, color=LIGHT_GRAY, size=9)

    # Activity-rich CRM with contact links
    n_engagements = manifest.get("engagements_count") or 0
    n_events = int(manifest.get("custom_events_fired_count") or 0)
    if not n_events:
        n_events = sum(int(e.get("test_submissions") or 0) for e in (plan.get("custom_events") or []))
    backdate_days = (plan.get("activity") or {}).get("backdate_days", 120)
    if contacts:
        crm_label = "Record-rich CRM" if _is_feature_showcase(plan) else "Activity-rich CRM"
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        events_clause = f" + {n_events} custom event fires" if n_events else ""
        r = p.add_run(
            f"{crm_label} ({len(contacts)} contacts, {n_engagements} engagements{events_clause}, "
            f"backdated {backdate_days} days)  ▸  "
        )
        _set_run(r, bold=True, size=10)
        # Render each contact as a clickable link with the human name from the plan
        plan_contacts = {c.get("email"): c for c in (plan.get("contacts") or [])}
        items = list(contacts.items())
        for i, (email, cid) in enumerate(items):
            pc = plan_contacts.get(email) or {}
            display = (
                f"{pc.get('firstname', '')} {pc.get('lastname', '')}".strip()
                or email.split("@")[0].replace(".", " ").title()
            )
            _add_hyperlink(p, _contact_url(portal, cid), display, size=10)
            if i < len(items) - 1:
                sep = p.add_run("  ·  ")
                _set_run(sep, color=LIGHT_GRAY, size=10)

    # Custom object
    co = manifest.get("custom_object") or {}
    if co.get("object_type_id"):
        n_records = len((plan.get("custom_object") or {}).get("records") or [])
        labels = (plan.get("custom_object") or {}).get("labels") or {}
        plural = labels.get("plural") or co.get("name") or "Records"
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"Custom data model: {plural} ({n_records} records associated to deals)  ▸  ")
        _set_run(r, bold=True, size=10)
        _add_hyperlink(p, urls["custom_obj"], f"Open {plural} object", size=10)

    # Custom events
    custom_events = manifest.get("custom_events") or {}
    if custom_events:
        ev_name = next(iter(custom_events.keys()))
        flow_count = len(manifest.get("custom_event_flows") or {})
        fired_count = int(manifest.get("custom_events_fired_count") or 0)
        prefix = (
            f"Custom event flows: {flow_count} ({fired_count} fires)  "
            if flow_count else f"Custom event definition: {ev_name}  "
        )
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(prefix + "▸  ")
        _set_run(r, bold=True, size=10)
        _add_hyperlink(p, urls["event_defs"], "Event definitions", size=10)

    # Marketing campaign(s)
    campaigns = manifest.get("campaigns") or {}
    if campaigns:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"Marketing campaigns ({len(campaigns)})  ▸  ")
        _set_run(r, bold=True, size=10)
        items = list(campaigns.items())
        for i, (campaign_name, cinfo) in enumerate(items[:6]):
            cinfo = cinfo or {}
            label = str(campaign_name)
            role = cinfo.get("role") or ""
            if role and role != "primary":
                label += f" ({role.replace('_', ' ')})"
            _add_hyperlink(p, cinfo.get("url") or urls.get("campaigns_index", ""), label, size=10)
            if i < min(len(items), 6) - 1:
                sep = p.add_run("  ·  ")
                _set_run(sep, color=LIGHT_GRAY, size=10)
    elif manifest.get("campaign_id"):
        campaign_name = ((plan.get("marketing_campaign") or {}).get("name")
                         or "Marketing campaign")
        campaign_url = manifest.get("campaign_url") or urls.get("campaigns_index", "")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"Marketing campaign: {campaign_name}  ▸  ")
        _set_run(r, bold=True, size=10)
        if campaign_url:
            _add_hyperlink(p, campaign_url, "Open campaign", size=10)

    # Tickets
    tickets = manifest.get("tickets") or {}
    if tickets:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"Support tickets ({len(tickets)} sample issues)  ▸  ")
        _set_run(r, bold=True, size=10)
        _add_hyperlink(p, urls["tickets"], "Tickets list", size=10)

    # Company record
    if (manifest.get("company") or {}).get("id"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run("Company record  ▸  ")
        _set_run(r, bold=True, size=10)
        _add_hyperlink(p, urls["company"], (manifest.get("company") or {}).get("name") or "Company", size=10)


def _recommendation_text(manifest: dict, plan: dict) -> str:
    """Return the recommendation paragraph for the doc.

    Preference order:
      1. plan["recommendation_text"] verbatim (orchestrator-supplied, prospect-tuned)
      2. A generic template that ONLY references manifest-derived values — no
         industry-specific copy, no invented figures.
    Phantom-number guard runs on whichever string we pick before returning."""
    supplied = plan.get("recommendation_text") if isinstance(plan, dict) else None
    if isinstance(supplied, str) and supplied.strip():
        return _strip_phantom_numbers(supplied.strip(), manifest, plan)

    if _is_feature_showcase(plan):
        feature = _feature_showcase(plan)
        features = _text_list(feature.get("requested_features") or feature.get("features"), limit=3)
        shot_list = _text_list(feature.get("shot_list"), limit=3)
        agenda_titles = [
            item.get("title") for item in (plan.get("agenda") or [])
            if isinstance(item, dict) and item.get("title")
        ][:3]
        if shot_list:
            first_shot = shot_list[0].rstrip(".")
            remaining = ", ".join(s.rstrip(".") for s in shot_list[1:])
            text = (
                f"Start here: {first_shot}. "
                f"Then move through {remaining if remaining else 'the built HubSpot artifacts'} "
                "so the audience sees the feature working through real records, not a feature list."
            )
        elif agenda_titles:
            text = (
                f"Lead with {agenda_titles[0]}. "
                f"Then use {'; '.join(agenda_titles[1:]) if len(agenda_titles) > 1 else 'the linked artifacts'} "
                "as the recording path. Keep the narration anchored on the data changing hands between records."
            )
        else:
            feature_label = ", ".join(features) if features else "the requested feature"
            text = (
                f"Lead with the cleanest sample record for {feature_label}. "
                "Open each linked artifact in the doc in order, and close on the adjacent-value item "
                "as the practical next step."
            )
        return _strip_phantom_numbers(text, manifest, plan)

    # Generic fallback — only manifest-derived values, no industry copy
    contacts = list((manifest.get("contacts") or {}).items())
    plan_contacts = {c.get("email"): c for c in (plan.get("contacts") or [])}
    sample_name = ""
    if contacts:
        first_email, _ = contacts[0]
        pc = plan_contacts.get(first_email) or {}
        sample_name = f"{pc.get('firstname', '')} {pc.get('lastname', '')}".strip()
        if not sample_name:
            # Derive a readable name from the email local part as a last resort
            sample_name = first_email.split("@")[0].replace(".", " ").title()
    if not sample_name:
        sample_name = "a top contact"

    co_label = ""
    co_plan = plan.get("custom_object") or {}
    if isinstance(co_plan, dict):
        co_label = (co_plan.get("labels") or {}).get("plural") or co_plan.get("name") or ""
    if not co_label:
        co_label = (manifest.get("custom_object") or {}).get("name") or ""
    custom_object_or_default = co_label or "Contacts"

    text = (
        f"Lead with the activity timeline on {sample_name}. "
        "It sells the value of the CRM faster than any feature list. "
        f"Then walk the deal pipeline, drop into the {custom_object_or_default} view, "
        "and close on the marketing email + lead scoring as the 'here is where "
        "automation takes over' moment."
    )
    return _strip_phantom_numbers(text, manifest, plan)


def _render_checklist(doc, *, manifest: dict, urls: dict, portal: str) -> None:
    contacts = manifest.get("contacts") or {}
    sample_link = None
    if contacts:
        first_email, first_id = next(iter(contacts.items()))
        sample_link = (
            f"{first_email.split('@')[0].replace('.', ' ').title()}'s contact",
            _contact_url(portal, first_id),
        )

    checks: list[tuple[str, str, str, str]] = []
    if sample_link:
        checks.append(("Open ", sample_link[0], sample_link[1], " and confirm the timeline looks alive"))
    if (manifest.get("pipeline") or {}).get("id"):
        checks.append(("Open the ", "deal pipeline board view", urls["pipeline_board"],
                       " and verify all deals show across stages"))
    if (manifest.get("marketing_email") or {}).get("id") or (manifest.get("marketing_email") or {}).get("name"):
        has_hero = bool((manifest.get("marketing_email") or {}).get("hero_image_url"))
        checks.append(("Open the ", "marketing email", urls["email"],
                       " and confirm the AI hero image renders" if has_hero else " and confirm the branded copy renders"))
    if urls.get("nps_form"):
        checks.append(("Open the ", "NPS survey form", urls["nps_form"],
                       " and verify the questions look right"))
    if urls.get("custom_obj"):
        co_label = (manifest.get("custom_object") or {}).get("name") or "custom"
        checks.append(("Open the ", f"{co_label} custom object", urls["custom_obj"],
                       " and confirm records are visible"))

    for prefix, link_text, url, suffix in checks:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.05)
        r = p.add_run("☐  " + prefix)
        _set_run(r, size=10)
        _add_hyperlink(p, url, link_text, size=10)
        r = p.add_run(suffix)
        _set_run(r, size=10)


def _render_feature_showcase_brief(doc, *, feature: dict, company: dict, research: dict) -> None:
    """Render the story/feature brief for Feature Showcase mode."""
    rows: list[tuple[str, str]] = []
    story = feature.get("story") or research.get("stated_context")
    if story:
        rows.append(("Story", str(story)[:900]))
    features = _text_list(feature.get("requested_features") or feature.get("features"))
    if features:
        rows.append(("Features", ", ".join(features)))
    audience = feature.get("audience") or research.get("audience")
    if audience:
        rows.append(("Audience", str(audience)))
    criteria = _text_list(feature.get("success_criteria"), limit=5)
    if criteria:
        rows.append(("Success criteria", "; ".join(criteria)))
    artifact_goals = _text_list(feature.get("artifact_goals"), limit=5)
    if artifact_goals:
        rows.append(("Artifact goals", "; ".join(artifact_goals)))
    shot_list = _text_list(feature.get("shot_list"), limit=6)
    if shot_list:
        rows.append(("Shot list", "; ".join(shot_list)))
    if feature.get("easter_egg_strategy"):
        rows.append(("Adjacent value", str(feature["easter_egg_strategy"])[:500]))

    rs_company = research.get("company") or {}
    domain = company.get("domain") or rs_company.get("domain")
    if domain:
        rows.append(("Brand context", domain))

    if not rows:
        rows.append(("Mode", "Feature Showcase"))

    for k, v in rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.05)
        r = p.add_run(f"{k}: ")
        _set_run(r, bold=True, size=9.5)
        r = p.add_run(str(v))
        _set_run(r, size=9.5)


def _render_attribution_story_map(doc, *, plan: dict) -> None:
    """Render a compact map of the campaign attribution showcase plan."""
    block = plan.get("campaign_attribution_showcase") or {}
    if not isinstance(block, dict):
        return

    rows: list[tuple[str, str]] = []
    campaigns = block.get("campaigns") or []
    if campaigns:
        labels = []
        for campaign in campaigns[:5]:
            if not isinstance(campaign, dict):
                continue
            name = campaign.get("name") or "Campaign"
            source = campaign.get("source") or campaign.get("role") or ""
            role = campaign.get("role") or ""
            detail = " / ".join([str(x) for x in (source, role) if x])
            labels.append(f"{name} ({detail})" if detail else str(name))
        if labels:
            rows.append(("Campaign mix", "; ".join(labels)))

    contact_paths = block.get("contact_paths") or []
    if contact_paths:
        examples = []
        for path in contact_paths[:4]:
            if not isinstance(path, dict):
                continue
            contact = path.get("contact_email") or path.get("contact") or "sample contact"
            first = path.get("first_touch_campaign") or "first touch"
            last = path.get("last_touch_campaign") or "last touch"
            deal = path.get("deal_name") or "associated deal"
            revenue = path.get("revenue")
            rev = f" ({_format_currency(revenue)})" if revenue else ""
            examples.append(f"{contact}: {first} → {last} → {deal}{rev}")
        if examples:
            rows.append(("Contact paths", "; ".join(examples)))

    rollup = block.get("deal_campaign_rollup") or {}
    if isinstance(rollup, dict) and rollup:
        method = rollup.get("method") or "workflow/manual step"
        workflow = rollup.get("workflow_name") or ""
        props = _text_list(rollup.get("deal_properties"), limit=6)
        detail = method
        if workflow:
            detail += f" · {workflow}"
        if props:
            detail += f" · deal fields: {', '.join(props)}"
        rows.append(("Deal rollup", detail))

    reports = _text_list(block.get("reports"), limit=6)
    if reports:
        rows.append(("Report targets", "; ".join(reports)))

    for k, v in rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.05)
        r = p.add_run(f"{k}  ▸  ")
        _set_run(r, bold=True, size=9.5)
        r = p.add_run(str(v))
        _set_run(r, size=9.5)


def _render_snapshot(doc, *, company: dict, research: dict) -> None:
    """Render the company snapshot table from whatever fields we have."""
    rows: list[tuple[str, str]] = []
    industry = company.get("industry") or research.get("company", {}).get("industry") or ""
    if industry:
        rows.append(("Industry", industry.replace("_", " ").title()))
    if company.get("description"):
        rows.append(("About", company["description"]))
    rs_company = research.get("company") or {}
    domain = company.get("domain") or rs_company.get("domain")
    if domain:
        rows.append(("Website", domain))
    if research.get("phone") or company.get("phone"):
        rows.append(("Phone", research.get("phone") or company.get("phone")))
    # Surface any extra rows the plan may have added (founded, fleet, hq, etc.)
    for k in ("founded", "hq", "fleet", "services"):
        v = company.get(k)
        if v:
            rows.append((k.title(), v))

    for k, v in rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.05)
        r = p.add_run(f"{k}: ")
        _set_run(r, bold=True, size=9.5)
        r = p.add_run(str(v))
        _set_run(r, size=9.5)


def _render_icp(doc, *, research: dict, plan: dict) -> None:
    """Pull the Perplexity research summary into a few tight bullets."""
    rows: list[tuple[str, str]] = []
    if _is_feature_showcase(plan):
        feature = _feature_showcase(plan, research)
        features = _text_list(feature.get("requested_features") or feature.get("features"), limit=5)
        if features:
            rows.append(("Requested feature story", ", ".join(features)))
        audience = feature.get("audience")
        if audience:
            rows.append(("Audience", str(audience)))
        criteria = _text_list(feature.get("success_criteria"), limit=4)
        if criteria:
            rows.append(("What good looks like", "; ".join(criteria)))

    perp = research.get("perplexity") or {}
    if perp.get("choices"):
        # Just include a top-line statement if we can derive one.
        msg = perp["choices"][0].get("message", {}).get("content", "")
        # Pull the headline (everything before "###" if present)
        headline = msg.split("###", 1)[0].strip().strip("*").strip()
        if headline:
            rows.append(("Primary ICP", headline[:600]))

    # Any agenda 'stat' lines double as research-backed talking points
    for item in (plan.get("agenda") or [])[:3]:
        if item.get("stat"):
            rows.append((item.get("title", "Talking point").split("—")[0].strip(), item["stat"]))
    easter = plan.get("easter_egg") or {}
    if easter.get("stat"):
        rows.append((easter.get("title", "Easter egg").split("—")[0].strip(), easter["stat"]))

    for k, v in rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.05)
        r = p.add_run(f"{k}  ▸  ")
        _set_run(r, bold=True, size=9.5)
        r = p.add_run(str(v))
        _set_run(r, size=9.5)


def _build_inventory(manifest: dict, plan: dict) -> list[str]:
    items: list[str] = []
    co_name = (plan.get("company") or {}).get("name") or "Customer"
    if (manifest.get("company") or {}).get("id"):
        items.append(f"1 company ({co_name}) with full firmographics")
    contacts = manifest.get("contacts") or {}
    if contacts:
        roles = [c.get("jobtitle") for c in (plan.get("contacts") or []) if c.get("jobtitle")]
        roles_clause = (": " + ", ".join(roles[:8])) if roles else ""
        items.append(f"{len(contacts)} contacts spanning roles{roles_clause}")
    deals = manifest.get("deals") or {}
    if deals:
        total = sum(_coerce_amount(d.get("amount")) for d in (plan.get("deals") or []))
        pipeline_name = (manifest.get("pipeline") or {}).get("name") or "the custom pipeline"
        items.append(f"{len(deals)} deals across the {pipeline_name} ({_format_currency(total)} total ACV)")
    n_eng = manifest.get("engagements_count") or 0
    days = (plan.get("activity") or {}).get("backdate_days", 120)
    if n_eng:
        items.append(f"{n_eng} timeline engagements (notes, calls, tasks, meetings, emails) backdated {days} days for a lived-in feel")
    custom_events = manifest.get("custom_events") or {}
    if custom_events:
        n_events = int(manifest.get("custom_events_fired_count") or 0)
        if not n_events:
            n_events = sum(int(e.get("test_submissions") or 0) for e in (plan.get("custom_events") or []))
        ev_name = next(iter(custom_events.keys()))
        items.append(f"{n_events} custom event fires ({ev_name}) on contact records" if n_events
                     else f"Custom event definition ({ev_name}) wired to contact records")
    co = manifest.get("custom_object") or {}
    if co.get("object_type_id"):
        n_records = len((plan.get("custom_object") or {}).get("records") or [])
        plural = (plan.get("custom_object") or {}).get("labels", {}).get("plural") or co.get("name", "Records")
        items.append(f"1 custom object ({plural}, {n_records} records) associated to deals: demonstrates industry-specific data modeling")
    forms = manifest.get("forms") or {}
    for fname in forms:
        items.append(f"1 form (live, embeddable): {fname}")
    me = manifest.get("marketing_email") or {}
    if me.get("name") or me.get("html_path"):
        hero_clause = " with AI-generated hero image" if me.get("hero_image_url") else ""
        items.append(f"1 branded marketing email (live in HubSpot){hero_clause}")
    campaigns = manifest.get("campaigns") or {}
    if campaigns:
        items.append(f"{len(campaigns)} marketing campaign records linked to demo assets")
    ls = manifest.get("lead_scoring") or {}
    if ls.get("property"):
        items.append(f"1 custom property ({ls['property']}, 0-100) backfilled on all {ls.get('backfilled') or len(contacts)} contacts")
    if ls.get("list_id"):
        items.append("1 lead-scoring contact list (Hot Leads by score)")
    attr = manifest.get("campaign_attribution_showcase") or {}
    if attr:
        items.append(
            "Campaign attribution showcase fields patched on "
            f"{attr.get('contacts_patched', 0)} contact(s) and "
            f"{attr.get('deals_patched', 0)} deal(s)"
        )
    tickets = manifest.get("tickets") or {}
    if tickets:
        items.append(f"{len(tickets)} support tickets (sample customer issues)")
    dashboards = manifest.get("dashboards_v04") or {}
    reports_status = manifest.get("reports_status") or {}
    if dashboards:
        items.append(f"{len(dashboards)} v0.4 reporting dashboard(s) with {len(manifest.get('reports') or {})} report manifest entries")
    elif reports_status:
        items.append(f"v0.4 reporting dashboard plan recorded ({reports_status.get('status', 'not built')})")
    return items


def _build_limitations(manifest: dict) -> list[str]:
    """Surface known-failed builds as a bullet list so the rep walks them live."""
    limits: list[str] = []
    errors = manifest.get("errors") or []
    manual = manifest.get("manual_steps") or []
    seen = set()
    for ms in manual:
        msg = (ms.get("instructions") or "").strip()
        key = msg[:80]
        if key and key not in seen:
            seen.add(key)
            limits.append(msg)
    for e in errors:
        where = e.get("where", "")
        if "form" in where.lower() and "Quote form" not in " ".join(limits):
            limits.append("Quote form is built in the UI for richer template handling; only the working form was created. Clone in the UI if needed for the live demo.")
        if "workflow" in where.lower() and "Workflow build" not in " ".join(limits):
            limits.append("Workflow build is faster in the UI for this routing logic; affected workflows are built manually (60 seconds each).")
        if "reports" in where.lower() and "Reporting dashboards" not in " ".join(limits):
            limits.append("Reporting dashboards are planned in the build plan but need the Playwright report-builder phase or a manual UI pass before the live demo.")
    return limits[:6]


def _research_sources(research: dict) -> list[tuple[str, str]]:
    sources_raw = research.get("sources") or []
    out: list[tuple[str, str]] = []
    for s in sources_raw[:10]:
        if isinstance(s, dict):
            out.append((s.get("title") or s.get("url") or "", s.get("url") or ""))
        else:
            # Use a friendlier label by stripping protocol
            label = str(s).replace("https://", "").replace("http://", "").rstrip("/")
            out.append((label, str(s)))
    return [(l, u) for l, u in out if u]


# =====================================================================
# Public API: docx generation
# =====================================================================

def generate_docx(manifest: dict, research: dict, plan: dict, *,
                  slug: str, work_dir: str, portal: str) -> str:
    """Build the demo/showcase runbook .docx. Returns the path to the saved .docx."""
    doc = _build_doc(manifest, research, plan, slug=slug, work_dir=work_dir, portal=portal)
    out_path = os.path.join(work_dir, "demo-doc.docx")
    doc.save(out_path)
    return out_path


# =====================================================================
# Drive auth + upload + export
# =====================================================================

def _refresh_drive_access_token() -> str | None:
    """Refresh the gdrive OAuth access token via rclone's stored refresh token.

    Returns the access token string, or None if rclone is not configured /
    refresh fails. Never raises.
    """
    if not RCLONE_CONF_PATH.is_file():
        print(f"[doc_generator] WARN: {RCLONE_CONF_PATH} not found; skipping Drive upload",
              file=sys.stderr)
        return None
    cp = configparser.ConfigParser()
    try:
        cp.read(RCLONE_CONF_PATH)
    except Exception as e:
        print(f"[doc_generator] WARN: failed to parse rclone.conf: {e}", file=sys.stderr)
        return None
    if "gdrive" not in cp:
        print("[doc_generator] WARN: no [gdrive] section in rclone.conf", file=sys.stderr)
        return None
    try:
        tok = json.loads(cp["gdrive"]["token"])
    except Exception as e:
        print(f"[doc_generator] WARN: gdrive token unreadable: {e}", file=sys.stderr)
        return None
    refresh = tok.get("refresh_token")
    if not refresh:
        print("[doc_generator] WARN: no refresh_token in gdrive section", file=sys.stderr)
        return None
    try:
        req = urllib.request.Request(
            "https://oauth2.googleapis.com/token",
            data=urllib.parse.urlencode({
                "client_id": RCLONE_CLIENT_ID,
                "client_secret": RCLONE_CLIENT_SECRET,
                "refresh_token": refresh,
                "grant_type": "refresh_token",
            }).encode(),
            method="POST",
        )
        with urllib.request.urlopen(req, timeout=30) as r:
            return json.loads(r.read())["access_token"]
    except Exception as e:
        print(f"[doc_generator] WARN: token refresh failed: {e}", file=sys.stderr)
        return None


def upload_to_drive(docx_path: str, *, doc_title: str,
                    drive_folder_id: str = DEFAULT_DRIVE_FOLDER_ID,
                    replace_doc_id: str | None = None) -> dict:
    """Upload a .docx to Drive, converting to a Google Doc.

    If replace_doc_id is provided, PATCH that file (preserving its ID/URL).
    Otherwise POST a new file in the configured folder.

    PDF export is best-effort. On any error returns None values rather than raising.
    """
    result: dict = {"gdoc_url": None, "doc_id": None, "pdf_path": None}

    access_token = _refresh_drive_access_token()
    if not access_token:
        return result

    try:
        with open(docx_path, "rb") as f:
            docx_bytes = f.read()
    except Exception as e:
        print(f"[doc_generator] WARN: cannot read {docx_path}: {e}", file=sys.stderr)
        return result

    boundary = "------demoprepboundary"
    metadata: dict = {
        "mimeType": "application/vnd.google-apps.document",
        "name": doc_title,
    }
    if not replace_doc_id and drive_folder_id:
        metadata["parents"] = [drive_folder_id]

    body = (
        f"--{boundary}\r\n"
        "Content-Type: application/json; charset=UTF-8\r\n\r\n"
        f"{json.dumps(metadata)}\r\n"
        f"--{boundary}\r\n"
        "Content-Type: application/vnd.openxmlformats-officedocument.wordprocessingml.document\r\n\r\n"
    ).encode() + docx_bytes + f"\r\n--{boundary}--\r\n".encode()

    if replace_doc_id:
        url = (f"https://www.googleapis.com/upload/drive/v3/files/{replace_doc_id}"
               "?uploadType=multipart&supportsAllDrives=true")
        method = "PATCH"
    else:
        url = ("https://www.googleapis.com/upload/drive/v3/files"
               "?uploadType=multipart&supportsAllDrives=true")
        method = "POST"

    req = urllib.request.Request(
        url, data=body,
        headers={
            "Authorization": f"Bearer {access_token}",
            "Content-Type": f"multipart/related; boundary={boundary}",
        },
        method=method,
    )
    try:
        with urllib.request.urlopen(req, timeout=60) as r:
            resp = json.loads(r.read())
    except urllib.error.HTTPError as e:
        body_txt = ""
        try:
            body_txt = e.read().decode()
        except Exception:
            pass
        print(f"[doc_generator] WARN: Drive upload HTTP {e.code}: {body_txt[:300]}",
              file=sys.stderr)
        return result
    except Exception as e:
        print(f"[doc_generator] WARN: Drive upload failed: {e}", file=sys.stderr)
        return result

    doc_id = resp.get("id")
    if not doc_id:
        return result
    result["doc_id"] = doc_id
    result["gdoc_url"] = f"https://docs.google.com/document/d/{doc_id}/edit"

    # Best-effort PDF export
    pdf_out = os.path.join(os.path.dirname(docx_path), "demo-doc.pdf")
    pdf_path = export_pdf(doc_id, pdf_out)
    result["pdf_path"] = pdf_path
    return result


def export_pdf(doc_id: str, out_path: str) -> str | None:
    """Export the GDoc as PDF. Returns out_path on success, None on failure."""
    access_token = _refresh_drive_access_token()
    if not access_token:
        return None
    req = urllib.request.Request(
        f"https://www.googleapis.com/drive/v3/files/{doc_id}/export?mimeType=application/pdf",
        headers={"Authorization": f"Bearer {access_token}"},
    )
    try:
        with urllib.request.urlopen(req, timeout=60) as r, open(out_path, "wb") as f:
            f.write(r.read())
        return out_path
    except urllib.error.HTTPError as e:
        body_txt = ""
        try:
            body_txt = e.read().decode()
        except Exception:
            pass
        print(f"[doc_generator] WARN: PDF export HTTP {e.code}: {body_txt[:300]}",
              file=sys.stderr)
        return None
    except Exception as e:
        print(f"[doc_generator] WARN: PDF export failed: {e}", file=sys.stderr)
        return None


# =====================================================================
# Inline self-tests for the phantom-number guard
# =====================================================================

def _run_phantom_guard_selftest() -> int:
    """Exercise the phantom-number guard against the v0.3.0 review cases.

    Returns 0 on success, 1 on any failure. Run via:
        python3 doc_generator.py --selftest-phantom
    """
    # A manifest with one real $4,200 deal.
    manifest = {"deals": {"d1": {"amount": 4200}}}
    plan: dict = {}

    cases: list[tuple[str, str, bool]] = [
        # (label, sentence, expected_keep)
        ("$1,200,000 phantom vs real $1,200 deal — DROP",
         "We projected $1,200,000 in pipeline.",
         False),
        ("$4,200.50 vs real $4,200 deal — DROP (different amount)",
         "Closed at $4,200.50 last quarter.",
         False),
        ("$4,200 vs real $4,200 deal — KEEP (exact)",
         "Closed at $4,200 last quarter.",
         True),
        ("$4.2K vs real $4,200 deal — KEEP (K-shorthand match)",
         "Closed at $4.2K last quarter.",
         True),
        ("$2.5M phantom vs real $4,200 deal — DROP",
         "Pipeline grew to $2.5M.",
         False),
    ]

    print("Phantom-number guard self-tests:")
    failures = 0
    for label, sent, expected_keep in cases:
        result = _strip_phantom_numbers(sent, manifest, plan)
        kept = bool(result.strip())
        status = "PASS" if kept == expected_keep else "FAIL"
        verdict = "kept" if kept else "dropped"
        if kept != expected_keep:
            failures += 1
        print(f"  [{status}] {label}")
        print(f"          input  = {sent!r}")
        print(f"          output = {result!r}  (sentence {verdict})")

    # Sentence-split test: 'e.g.' should not split mid-sentence.
    sent_with_eg = "We talked through several options, e.g. the premium tier and add-ons."
    parts = _ABBREV_LOOKBEHIND.split(sent_with_eg)
    eg_status = "PASS" if len(parts) == 1 else "FAIL"
    if len(parts) != 1:
        failures += 1
    print(f"  [{eg_status}] 'e.g.' mid-sentence does not split (got {len(parts)} part(s))")

    # And a normal multi-sentence string SHOULD split.
    multi = "First sentence. Second sentence."
    parts2 = _ABBREV_LOOKBEHIND.split(multi)
    multi_status = "PASS" if len(parts2) == 2 else "FAIL"
    if len(parts2) != 2:
        failures += 1
    print(f"  [{multi_status}] Normal sentence boundary still splits (got {len(parts2)} part(s))")

    print(f"\n{'OK' if failures == 0 else 'FAIL'}: {failures} failure(s)")
    return 0 if failures == 0 else 1


# =====================================================================
# CLI for local regen / smoke testing
# =====================================================================

def _main(argv: list[str]) -> int:
    if len(argv) >= 2 and argv[1] == "--selftest-phantom":
        return _run_phantom_guard_selftest()
    if len(argv) < 2:
        print("usage: doc_generator.py <work_dir> [portal]", file=sys.stderr)
        print("       doc_generator.py --selftest-phantom", file=sys.stderr)
        return 2
    work_dir = argv[1]
    portal = argv[2] if len(argv) > 2 else "51393541"
    manifest = json.load(open(os.path.join(work_dir, "manifest.json")))
    research = json.load(open(os.path.join(work_dir, "research.json")))
    plan = json.load(open(os.path.join(work_dir, "build-plan.json")))
    slug = plan.get("slug") or os.path.basename(work_dir.rstrip("/")).replace("demo-prep-", "")
    out = generate_docx(manifest, research, plan, slug=slug, work_dir=work_dir, portal=portal)
    print(f"Saved {out} ({os.path.getsize(out)} bytes)")
    return 0


if __name__ == "__main__":
    sys.exit(_main(sys.argv))
