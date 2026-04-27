#!/usr/bin/env python3
"""Generate the per-prospect demo runbook .docx, then (optionally) upload to Google Drive.

Ported from /tmp/demo-prep-shipperzinc/{make-doc.py, update-doc.py, export-pdf.py}.

The visual layout (banner, status pills, agenda block, easter egg, "also built",
recommendation, page-2 supporting documentation) is identical to the locked
Shipperz runbook. All slug-specific copy and IDs come from build-plan.json,
manifest.json, and research.json so the same generator works for any prospect.

Public API:
    generate_docx(manifest, research, plan, *, slug, work_dir, portal) -> str
    upload_to_drive(docx_path, *, doc_title, drive_folder_id=..., replace_doc_id=None) -> dict
    export_pdf(doc_id, out_path) -> str | None
"""
from __future__ import annotations

import configparser
import datetime
import json
import os
import pathlib
import sys
import urllib.error
import urllib.parse
import urllib.request

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# ---- Brand palette (kept identical to make-doc.py) ----
SHIPPERZ_ORANGE = RGBColor(0xFF, 0x6B, 0x35)  # default accent / kept name for parity
SHIPPERZ_DARK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x55, 0x55, 0x55)
LIGHT_GRAY = RGBColor(0x99, 0x99, 0x99)
SUCCESS_GREEN = RGBColor(0x2C, 0x84, 0x4F)
WARN_AMBER = RGBColor(0xB1, 0x6E, 0x05)
NOT_BUILT_RED = RGBColor(0xB9, 0x1C, 0x1C)
BLUE = RGBColor(0x06, 0x6E, 0xA0)


# ---- Drive / OAuth defaults ----
DEFAULT_DRIVE_FOLDER_ID = "1SzHT9uhFUUcFIAh5z2LVCAq2Wt0OADjY"
RCLONE_CONF_PATH = pathlib.Path.home() / ".config/rclone/rclone.conf"
# Same client_id/secret used by rclone's gdrive backend (matches update-doc.py).
RCLONE_CLIENT_ID = "202264815644.apps.googleusercontent.com"
RCLONE_CLIENT_SECRET = "X4Z3ca8xfWDb1Voo-F9a7ZxJ"


# =====================================================================
# docx helpers (verbatim from make-doc.py)
# =====================================================================

def _set_run(run, *, color=None, bold=False, italic=False, size=None, font="Arial"):
    run.font.name = font
    if color is not None:
        run.font.color.rgb = color
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size is not None:
        run.font.size = Pt(size)


def _add_hyperlink(paragraph, url, text, *, color=BLUE, underline=True, bold=False, size=10):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rPr.append(sz)
    if color is not None:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), "%02X%02X%02X" % (color[0], color[1], color[2]))
        rPr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
    if bold:
        rPr.append(OxmlElement("w:b"))
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _shade_paragraph(p, fill_hex, left_border_hex=None):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)
    if left_border_hex:
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "24")
        left.set(qn("w:space"), "4")
        left.set(qn("w:color"), left_border_hex)
        pBdr.append(left)
        pPr.append(pBdr)


def _bottom_border(p, color_hex="FF6B35"):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _h2(doc, text, *, color=SHIPPERZ_ORANGE, before=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text.upper())
    _set_run(r, color=color, bold=True, size=10)


def _status_pill(p, text, color):
    r = p.add_run(f"[{text}] ")
    _set_run(r, color=color, bold=True, size=9)


# =====================================================================
# URL builders
# =====================================================================

def _hub_urls(portal: str, manifest: dict) -> dict:
    """Build the table of HubSpot deep links from portal id + manifest."""
    pipeline_id = (manifest.get("pipeline") or {}).get("id", "")
    company_id = (manifest.get("company") or {}).get("id", "")
    custom_obj_id = (manifest.get("custom_object") or {}).get("object_type_id", "")
    forms = manifest.get("forms") or {}
    nps_form_id = ""
    for fname, fid in forms.items():
        if "nps" in fname.lower():
            nps_form_id = fid
            break
    if not nps_form_id and forms:
        nps_form_id = next(iter(forms.values()))
    email_id = (manifest.get("marketing_email") or {}).get("id", "")
    lead_score_prop = (manifest.get("lead_scoring") or {}).get("property", "demo_lead_score")
    return {
        "pipeline_board": f"https://app.hubspot.com/contacts/{portal}/objects/0-3/views/all/board?pipeline={pipeline_id}",
        "pipeline_list": f"https://app.hubspot.com/contacts/{portal}/objects/0-3/views/all/list?pipeline={pipeline_id}",
        "company": f"https://app.hubspot.com/contacts/{portal}/record/0-2/{company_id}",
        "custom_obj": f"https://app.hubspot.com/contacts/{portal}/objects/{custom_obj_id}" if custom_obj_id else "",
        "nps_form": f"https://app.hubspot.com/forms/{portal}/editor/{nps_form_id}/edit/form" if nps_form_id else "",
        "email": f"https://app.hubspot.com/email/{portal}/edit/{email_id}/edit/content" if email_id else f"https://app.hubspot.com/email/{portal}/manage/state/all",
        "workflows": f"https://app.hubspot.com/workflows/{portal}/view/all-workflows",
        "forms_index": f"https://app.hubspot.com/forms/{portal}",
        "landing_pages": f"https://app.hubspot.com/website/{portal}/landing-pages",
        "contacts_list": f"https://app.hubspot.com/contacts/{portal}/objects/0-1/views/all/list",
        "lead_score_prop": f"https://app.hubspot.com/property-settings/{portal}/properties?type=0-1&property={lead_score_prop}",
        "tickets": f"https://app.hubspot.com/contacts/{portal}/objects/0-5/views/all/list",
        "event_defs": f"https://app.hubspot.com/events/{portal}/manage/event-definitions",
    }


def _deal_url(portal: str, did: str) -> str:
    return f"https://app.hubspot.com/contacts/{portal}/record/0-3/{did}"


def _contact_url(portal: str, cid: str) -> str:
    return f"https://app.hubspot.com/contacts/{portal}/record/0-1/{cid}"


# =====================================================================
# Banner discovery
# =====================================================================

def _banner_path(work_dir: str, slug: str) -> str | None:
    """Return the first banner image that exists, or None."""
    candidates = [
        os.path.join(work_dir, f"{slug}-banner.png"),
        os.path.join(work_dir, "shipperz-banner.png") if slug == "shipperzinc" else None,
        os.path.join(work_dir, "hero-image-email.png"),
        os.path.join(work_dir, "hero-image.png"),
    ]
    for c in candidates:
        if c and os.path.isfile(c):
            return c
    return None


# =====================================================================
# Doc body
# =====================================================================

def _format_currency(amount) -> str:
    try:
        return f"${int(round(float(amount))):,}"
    except (TypeError, ValueError):
        return str(amount or "")


def _build_doc(manifest: dict, research: dict, plan: dict, *,
               slug: str, work_dir: str, portal: str) -> Document:
    company = plan.get("company") or {}
    company_name = company.get("name") or (manifest.get("company") or {}).get("name") or "Customer"
    urls = _hub_urls(portal, manifest)
    agenda = plan.get("agenda") or []
    easter = plan.get("easter_egg") or {}

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.45)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(1)
    style.paragraph_format.line_spacing = 1.10

    # ---------------- PAGE 1 ----------------

    # Banner
    banner = _banner_path(work_dir, slug)
    if banner:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        try:
            run.add_picture(banner, width=Inches(5.6))
        except Exception:
            # If the image is malformed, skip the banner rather than crash.
            pass

    # Title bar
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("HubSpot Demo Prep")
    _set_run(r, color=SHIPPERZ_DARK, bold=True, size=18)
    _bottom_border(p)

    # Subtitle line
    today = datetime.date.today().strftime("%B %-d, %Y")
    location = company.get("location") or company.get("hq") or ""
    subtitle = f"{today}   ·   Sandbox {portal}   ·   Prepared for: {company_name}"
    if location:
        subtitle += f" ({location})"
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(subtitle)
    _set_run(r, color=GRAY, size=9)

    # Intro paragraph (rep input + what was built)
    rep_input = (research.get("stated_context") or "").strip()
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    if rep_input:
        r = p.add_run("Rep input: ")
        _set_run(r, bold=True, size=9.5)
        r = p.add_run(rep_input + "  ")
        _set_run(r, italic=True, color=GRAY, size=9.5)
    r = p.add_run("Built: ")
    _set_run(r, bold=True, size=9.5)
    built_summary = _built_summary(manifest, plan)
    r = p.add_run(built_summary)
    _set_run(r, size=9.5)

    # ---- AGENDA ----
    _h2(doc, "Agenda (from sales rep)", before=8)
    for i, item in enumerate(agenda, 1):
        _render_agenda_item(doc, i, item, manifest=manifest, urls=urls, portal=portal)

    # ---- EASTER EGG ----
    if easter:
        _render_easter_egg(doc, easter, manifest=manifest, urls=urls, portal=portal)

    # ---- ALSO BUILT ----
    _h2(doc, "Also built (pull these into the demo if useful)", before=8)
    _render_also_built(doc, manifest=manifest, plan=plan, urls=urls, portal=portal)

    # ---- RECOMMENDATION ----
    _h2(doc, "Recommendation", before=6)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(_recommendation_text(manifest, plan))
    _set_run(r, size=9.5)

    # ---------------- PAGE 2 ----------------

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Supporting documentation")
    _set_run(r, color=SHIPPERZ_DARK, bold=True, size=16)
    _bottom_border(p)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Reference material for the demo. Page 1 is the runbook; this page is your context cheat-sheet.")
    _set_run(r, color=GRAY, size=9)

    # Pre-demo checklist
    _h2(doc, "Pre-demo checklist", before=2)
    _render_checklist(doc, manifest=manifest, urls=urls, portal=portal)

    # Company snapshot
    _h2(doc, f"{company_name} snapshot", before=8)
    _render_snapshot(doc, company=company, research=research)

    # ICP / research
    _h2(doc, "ICP and pain-point research (Perplexity, sonar)", before=8)
    _render_icp(doc, research=research, plan=plan)

    # Full inventory
    _h2(doc, f"Full build inventory (sandbox {portal})", before=8)
    for item in _build_inventory(manifest, plan):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.18)
        r = p.add_run(f"•  {item}")
        _set_run(r, size=9.5)

    # Limitations
    limits = _build_limitations(manifest)
    if limits:
        _h2(doc, "Known build limitations", before=8)
        for item in limits:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Inches(0.18)
            r = p.add_run(f"•  {item}")
            _set_run(r, color=GRAY, size=9.5)

    # Sources
    sources = _research_sources(research)
    if sources:
        _h2(doc, "Sources", before=8)
        for label, url in sources:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Inches(0.18)
            _add_hyperlink(p, url, label, size=9.5)

    # Cleanup
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("When done with this demo: ")
    _set_run(r, color=GRAY, italic=True, size=9)
    mono = p.add_run(f"python3 ~/.claude/skills/hubspot-demo-prep/builder.py cleanup {slug}")
    mono.font.name = "Menlo"
    mono.font.size = Pt(8)
    mono.font.color.rgb = GRAY

    return doc


# =====================================================================
# Section renderers
# =====================================================================

def _built_summary(manifest: dict, plan: dict) -> str:
    parts = []
    n_contacts = len(manifest.get("contacts") or {})
    n_engagements = manifest.get("engagements_count") or 0
    n_deals = len(manifest.get("deals") or {})
    if n_contacts:
        parts.append(f"{n_contacts} contacts")
    if n_engagements:
        parts.append(f"{n_engagements} timeline events")
    if n_deals:
        parts.append(f"a {n_deals}-deal pipeline")
    if (manifest.get("custom_object") or {}).get("name"):
        parts.append(f"a custom {manifest['custom_object']['name']} object")
    if manifest.get("forms"):
        parts.append("a live NPS survey" if any("nps" in n.lower() for n in manifest["forms"]) else "a live form")
    if (manifest.get("marketing_email") or {}).get("name") or (manifest.get("marketing_email") or {}).get("html_path"):
        parts.append("a branded marketing email with AI hero image")
    if (manifest.get("lead_scoring") or {}).get("property"):
        backfilled = manifest["lead_scoring"].get("backfilled") or n_contacts
        parts.append(f"lead scoring on all {backfilled} contacts")
    company_name = (plan.get("company") or {}).get("name") or "the prospect"
    summary = f"activity-rich {company_name} portal with " + ", ".join(parts) + "."
    if manifest.get("manual_steps"):
        n_workflow_manual = sum(1 for ms in manifest["manual_steps"]
                                if "workflow" in (ms.get("item") or "").lower())
        if n_workflow_manual:
            summary += (f" {n_workflow_manual} workflow step(s) hit API limits and need a "
                        "60-second live build (called out per agenda item below).")
    return summary


def _agenda_status_lines(item: dict, idx: int, manifest: dict, urls: dict, portal: str) -> list[tuple[str, RGBColor, str, str]]:
    """For each agenda item produce a list of (pill_label, pill_color, body_text, link_url[, link_text]).

    Falls back to a generic 'open in HubSpot' link from agenda.show_label.
    Returns list of tuples: (pill, color, prefix, link_url, link_label)
    """
    lines: list[tuple[str, RGBColor, str, str, str]] = []
    title_l = (item.get("title") or "").lower()

    forms = manifest.get("forms") or {}
    has_nps = any("nps" in n.lower() for n in forms)
    has_quote_form = any("quote" in n.lower() for n in forms)
    has_email = bool((manifest.get("marketing_email") or {}).get("name")
                     or (manifest.get("marketing_email") or {}).get("html_path"))
    has_workflows = bool(manifest.get("workflows"))
    has_landing = bool(manifest.get("landing_page"))

    if "nurtur" in title_l or "drip" in title_l or "follow-up" in title_l:
        if has_email:
            lines.append(("BUILT", SUCCESS_GREEN,
                          "Branded marketing email with AI hero image  ▸  ",
                          urls["email"], "Open email in HubSpot"))
        if not has_workflows:
            lines.append(("BUILD LIVE", WARN_AMBER,
                          "Workflow itself (HubSpot v4 flows API limit)  ▸  ",
                          urls["workflows"], "Workflows → Create new"))
    elif "landing" in title_l or "inbound" in title_l or "form" in title_l:
        if not (has_quote_form and has_landing):
            lines.append(("NOT BUILT", NOT_BUILT_RED,
                          "Quote form rejected by HubSpot Forms API; landing page is a Marketing Hub Pro+ feature. Walk this one live.  ▸  ",
                          urls["forms_index"], "Forms"))
            lines.append(("ANALOG", BLUE,
                          ("The NPS form (built) is a working example of how the Quote form would behave once rebuilt in the UI  ▸  "
                           if has_nps else "Use the live form as an analog when walking through this  ▸  "),
                          urls["nps_form"] or urls["forms_index"], "Open NPS form" if has_nps else "Forms"))
        else:
            lines.append(("BUILT", SUCCESS_GREEN,
                          "Quote form + landing page live  ▸  ",
                          urls["forms_index"], "Open form"))
    elif "nps" in title_l or "feedback" in title_l or "survey" in title_l:
        if has_nps:
            lines.append(("BUILT", SUCCESS_GREEN,
                          f"{(plan_company_label(manifest))} NPS Survey form (live)  ▸  ".replace("  ", " "),
                          urls["nps_form"] or urls["forms_index"], "Open form"))
        if not has_workflows:
            lines.append(("BUILD LIVE", WARN_AMBER,
                          "Routing workflow (same v4 flows API limit)  ▸  ",
                          urls["workflows"], "Workflows → Create new"))
    else:
        # Generic fallback — single line with whatever URL we can derive
        link_label = item.get("show_label") or "Open in HubSpot"
        link_url = urls.get("workflows") or ""
        lines.append(("BUILT" if has_workflows else "BUILD LIVE",
                      SUCCESS_GREEN if has_workflows else WARN_AMBER,
                      "  ▸  ", link_url, link_label))
    return lines


def plan_company_label(manifest: dict) -> str:
    return (manifest.get("company") or {}).get("name") or ""


def _render_agenda_item(doc, idx: int, item: dict, *, manifest: dict, urls: dict, portal: str) -> None:
    title = item.get("title") or ""
    why = item.get("why") or ""

    # Heading
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{idx}.  ")
    _set_run(r, color=SHIPPERZ_ORANGE, bold=True, size=10.5)
    r = p.add_run(title)
    _set_run(r, color=SHIPPERZ_DARK, bold=True, size=10.5)

    # Why / description
    if why:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.22)
        r = p.add_run(why)
        _set_run(r, color=GRAY, italic=True, size=9.5)

    # Status lines
    for pill, color, body, link_url, link_label in _agenda_status_lines(item, idx, manifest, urls, portal):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.22)
        _status_pill(p, pill, color)
        r = p.add_run(body)
        _set_run(r, size=9.5)
        if link_url:
            _add_hyperlink(p, link_url, link_label, size=9.5)


def _render_easter_egg(doc, easter: dict, *, manifest: dict, urls: dict, portal: str) -> None:
    title = easter.get("title") or ""
    why = easter.get("why") or ""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(0)
    _shade_paragraph(p, "FFF3EC", "FF6B35")
    r = p.add_run(f"★  EASTER EGG  ·  {title}")
    _set_run(r, color=SHIPPERZ_ORANGE, bold=True, size=10)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    _shade_paragraph(p, "FFF3EC", "FF6B35")
    _status_pill(p, "BUILT", SUCCESS_GREEN)
    if why:
        r = p.add_run(f"{why}  ")
        _set_run(r, color=SHIPPERZ_DARK, italic=True, size=9.5)
    _add_hyperlink(p, urls["contacts_list"], "Sortable contact list", size=9.5)
    sep = p.add_run("  ·  ")
    _set_run(sep, color=LIGHT_GRAY, size=9.5)
    # First contact as a sample
    contacts = manifest.get("contacts") or {}
    if contacts:
        first_email, first_id = next(iter(contacts.items()))
        sample_label = "Sample: " + first_email.split("@")[0].replace(".", " ").title()
        _add_hyperlink(p, _contact_url(portal, first_id), sample_label, size=9.5)
        sep = p.add_run("  ·  ")
        _set_run(sep, color=LIGHT_GRAY, size=9.5)
    _add_hyperlink(p, urls["lead_score_prop"], "Property settings", size=9.5)


def _render_also_built(doc, *, manifest: dict, plan: dict, urls: dict, portal: str) -> None:
    deals = manifest.get("deals") or {}
    contacts = manifest.get("contacts") or {}

    # Pipeline + deals
    if (manifest.get("pipeline") or {}).get("id"):
        total = sum((d.get("amount") or 0) for d in (plan.get("deals") or []))
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"Active deal pipeline ({len(deals)} deals, {_format_currency(total)} total)  ▸  ")
        _set_run(r, bold=True, size=10)
        _add_hyperlink(p, urls["pipeline_board"], "Board view", size=10)
        sep = p.add_run("  ·  ")
        _set_run(sep, color=LIGHT_GRAY, size=10)
        _add_hyperlink(p, urls["pipeline_list"], "List view", size=10)

        # Each deal individually (with stage label from plan)
        plan_deal_lookup = {d.get("name"): d for d in (plan.get("deals") or [])}
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.22)
        items = list(deals.items())
        for i, (name, did) in enumerate(items):
            pd = plan_deal_lookup.get(name) or {}
            stage = pd.get("stage") or ""
            amount = _format_currency(pd.get("amount"))
            label = f"{stage}: {name} ({amount})" if stage else f"{name} ({amount})"
            _add_hyperlink(p, _deal_url(portal, did), label, size=9)
            if i < len(items) - 1:
                sep = p.add_run("  ·  ")
                _set_run(sep, color=LIGHT_GRAY, size=9)

    # Activity-rich CRM with contact links
    n_engagements = manifest.get("engagements_count") or 0
    n_events = sum(int(e.get("test_submissions") or 0) for e in (plan.get("custom_events") or []))
    backdate_days = (plan.get("activity") or {}).get("backdate_days", 120)
    if contacts:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        events_clause = f" + {n_events} custom event fires" if n_events else ""
        r = p.add_run(
            f"Activity-rich CRM ({len(contacts)} contacts, {n_engagements} engagements{events_clause}, "
            f"backdated {backdate_days} days)  ▸  "
        )
        _set_run(r, bold=True, size=10)
        # Render each contact as a clickable link with the human name from the plan
        plan_contacts = {c.get("email"): c for c in (plan.get("contacts") or [])}
        items = list(contacts.items())
        for i, (email, cid) in enumerate(items):
            pc = plan_contacts.get(email) or {}
            display = (
                f"{pc.get('firstname', '')} {pc.get('lastname', '')}".strip()
                or email.split("@")[0].replace(".", " ").title()
            )
            _add_hyperlink(p, _contact_url(portal, cid), display, size=10)
            if i < len(items) - 1:
                sep = p.add_run("  ·  ")
                _set_run(sep, color=LIGHT_GRAY, size=10)

    # Custom object
    co = manifest.get("custom_object") or {}
    if co.get("object_type_id"):
        n_records = len((plan.get("custom_object") or {}).get("records") or [])
        labels = (plan.get("custom_object") or {}).get("labels") or {}
        plural = labels.get("plural") or co.get("name") or "Records"
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"Custom data model: {plural} ({n_records} records associated to deals)  ▸  ")
        _set_run(r, bold=True, size=10)
        _add_hyperlink(p, urls["custom_obj"], f"Open {plural} object", size=10)

    # Custom events
    custom_events = manifest.get("custom_events") or {}
    if custom_events:
        ev_name = next(iter(custom_events.keys()))
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"Custom event definition: {ev_name}  ▸  ")
        _set_run(r, bold=True, size=10)
        _add_hyperlink(p, urls["event_defs"], "Event definitions", size=10)

    # Tickets
    tickets = manifest.get("tickets") or {}
    if tickets:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"Support tickets ({len(tickets)} sample issues)  ▸  ")
        _set_run(r, bold=True, size=10)
        _add_hyperlink(p, urls["tickets"], "Tickets list", size=10)

    # Company record
    if (manifest.get("company") or {}).get("id"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run("Company record  ▸  ")
        _set_run(r, bold=True, size=10)
        _add_hyperlink(p, urls["company"], (manifest.get("company") or {}).get("name") or "Company", size=10)


def _recommendation_text(manifest: dict, plan: dict) -> str:
    contacts = list((manifest.get("contacts") or {}).items())
    plan_contacts = {c.get("email"): c for c in (plan.get("contacts") or [])}
    sample_name = ""
    if contacts:
        first_email, _ = contacts[0]
        pc = plan_contacts.get(first_email) or {}
        sample_name = f"{pc.get('firstname', '')} {pc.get('lastname', '')}".strip()
    co_name = (plan.get("custom_object") or {}).get("labels", {}).get("plural") or "the custom object"

    if sample_name:
        intro = f"Lead with the activity timeline on {sample_name}."
    else:
        intro = "Lead with the activity timeline on a top contact."
    return (
        f"{intro} It sells the value of the CRM faster than any feature list. "
        f"Then walk the deal pipeline, drop into the {co_name} object to show industry-specific "
        "data modeling, and close on the marketing email + lead scoring as the “here's where "
        "automation takes over” moment. Build the lead-nurture workflow live in 60 seconds to "
        "show how easy it is for a no-marketing-team setup."
    )


def _render_checklist(doc, *, manifest: dict, urls: dict, portal: str) -> None:
    contacts = manifest.get("contacts") or {}
    sample_link = None
    if contacts:
        first_email, first_id = next(iter(contacts.items()))
        sample_link = (
            f"{first_email.split('@')[0].replace('.', ' ').title()}'s contact",
            _contact_url(portal, first_id),
        )

    checks: list[tuple[str, str, str, str]] = []
    if sample_link:
        checks.append(("Open ", sample_link[0], sample_link[1], " and confirm the timeline looks alive"))
    if (manifest.get("pipeline") or {}).get("id"):
        checks.append(("Open the ", "deal pipeline board view", urls["pipeline_board"],
                       " and verify all deals show across stages"))
    if (manifest.get("marketing_email") or {}).get("id") or (manifest.get("marketing_email") or {}).get("name"):
        checks.append(("Open the ", "marketing email", urls["email"],
                       " and confirm the AI hero image renders"))
    if urls.get("nps_form"):
        checks.append(("Open the ", "NPS survey form", urls["nps_form"],
                       " and verify the questions look right"))
    if urls.get("custom_obj"):
        co_label = (manifest.get("custom_object") or {}).get("name") or "custom"
        checks.append(("Open the ", f"{co_label} custom object", urls["custom_obj"],
                       " and confirm records are visible"))

    for prefix, link_text, url, suffix in checks:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.05)
        r = p.add_run("☐  " + prefix)
        _set_run(r, size=10)
        _add_hyperlink(p, url, link_text, size=10)
        r = p.add_run(suffix)
        _set_run(r, size=10)


def _render_snapshot(doc, *, company: dict, research: dict) -> None:
    """Render the company snapshot table from whatever fields we have."""
    rows: list[tuple[str, str]] = []
    industry = company.get("industry") or research.get("company", {}).get("industry") or ""
    if industry:
        rows.append(("Industry", industry.replace("_", " ").title()))
    if company.get("description"):
        rows.append(("About", company["description"]))
    rs_company = research.get("company") or {}
    domain = company.get("domain") or rs_company.get("domain")
    if domain:
        rows.append(("Website", domain))
    if research.get("phone") or company.get("phone"):
        rows.append(("Phone", research.get("phone") or company.get("phone")))
    # Surface any extra rows the plan may have added (founded, fleet, hq, etc.)
    for k in ("founded", "hq", "fleet", "services"):
        v = company.get(k)
        if v:
            rows.append((k.title(), v))

    for k, v in rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.05)
        r = p.add_run(f"{k}: ")
        _set_run(r, bold=True, size=9.5)
        r = p.add_run(str(v))
        _set_run(r, size=9.5)


def _render_icp(doc, *, research: dict, plan: dict) -> None:
    """Pull the Perplexity research summary into a few tight bullets."""
    rows: list[tuple[str, str]] = []
    perp = research.get("perplexity") or {}
    if perp.get("choices"):
        # Just include a top-line statement if we can derive one.
        msg = perp["choices"][0].get("message", {}).get("content", "")
        # Pull the headline (everything before "###" if present)
        headline = msg.split("###", 1)[0].strip().strip("*").strip()
        if headline:
            rows.append(("Primary ICP", headline[:600]))

    # Any agenda 'stat' lines double as research-backed talking points
    for item in (plan.get("agenda") or [])[:3]:
        if item.get("stat"):
            rows.append((item.get("title", "Talking point").split("—")[0].strip(), item["stat"]))
    easter = plan.get("easter_egg") or {}
    if easter.get("stat"):
        rows.append((easter.get("title", "Easter egg").split("—")[0].strip(), easter["stat"]))

    for k, v in rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.05)
        r = p.add_run(f"{k}  ▸  ")
        _set_run(r, bold=True, size=9.5)
        r = p.add_run(str(v))
        _set_run(r, size=9.5)


def _build_inventory(manifest: dict, plan: dict) -> list[str]:
    items: list[str] = []
    co_name = (plan.get("company") or {}).get("name") or "Customer"
    if (manifest.get("company") or {}).get("id"):
        items.append(f"1 company ({co_name}) with full firmographics")
    contacts = manifest.get("contacts") or {}
    if contacts:
        roles = [c.get("jobtitle") for c in (plan.get("contacts") or []) if c.get("jobtitle")]
        roles_clause = (": " + ", ".join(roles[:8])) if roles else ""
        items.append(f"{len(contacts)} contacts spanning roles{roles_clause}")
    deals = manifest.get("deals") or {}
    if deals:
        total = sum((d.get("amount") or 0) for d in (plan.get("deals") or []))
        pipeline_name = (manifest.get("pipeline") or {}).get("name") or "the custom pipeline"
        items.append(f"{len(deals)} deals across the {pipeline_name} ({_format_currency(total)} total ACV)")
    n_eng = manifest.get("engagements_count") or 0
    days = (plan.get("activity") or {}).get("backdate_days", 120)
    if n_eng:
        items.append(f"{n_eng} timeline engagements (notes, calls, tasks, meetings, emails) backdated {days} days for a lived-in feel")
    custom_events = manifest.get("custom_events") or {}
    if custom_events:
        n_events = sum(int(e.get("test_submissions") or 0) for e in (plan.get("custom_events") or []))
        ev_name = next(iter(custom_events.keys()))
        items.append(f"{n_events} custom event fires ({ev_name}) on contact records" if n_events
                     else f"Custom event definition ({ev_name}) wired to contact records")
    co = manifest.get("custom_object") or {}
    if co.get("object_type_id"):
        n_records = len((plan.get("custom_object") or {}).get("records") or [])
        plural = (plan.get("custom_object") or {}).get("labels", {}).get("plural") or co.get("name", "Records")
        items.append(f"1 custom object ({plural}, {n_records} records) associated to deals: demonstrates industry-specific data modeling")
    forms = manifest.get("forms") or {}
    for fname in forms:
        items.append(f"1 form (live, embeddable): {fname}")
    me = manifest.get("marketing_email") or {}
    if me.get("name") or me.get("html_path"):
        items.append("1 branded marketing email (live in HubSpot) with AI-generated hero image")
    ls = manifest.get("lead_scoring") or {}
    if ls.get("property"):
        items.append(f"1 custom property ({ls['property']}, 0-100) backfilled on all {ls.get('backfilled') or len(contacts)} contacts")
    if ls.get("list_id"):
        items.append("1 lead-scoring contact list (Hot Leads by score)")
    tickets = manifest.get("tickets") or {}
    if tickets:
        items.append(f"{len(tickets)} support tickets (sample customer issues)")
    return items


def _build_limitations(manifest: dict) -> list[str]:
    """Surface known-failed builds as a bullet list so the rep walks them live."""
    limits: list[str] = []
    errors = manifest.get("errors") or []
    manual = manifest.get("manual_steps") or []
    seen = set()
    for ms in manual:
        msg = (ms.get("instructions") or "").strip()
        key = msg[:80]
        if key and key not in seen:
            seen.add(key)
            limits.append(msg)
    for e in errors:
        where = e.get("where", "")
        if "form" in where.lower() and "Quote form" not in " ".join(limits):
            limits.append("Quote form returned an error from the Forms API; only the working form was created. Clone in the UI if needed for the live demo.")
        if "workflow" in where.lower() and "Workflow v4" not in " ".join(limits):
            limits.append("Workflow v4 flows API rejected our payload shape; affected workflows must be created in the UI (60 seconds each).")
    return limits[:6]


def _research_sources(research: dict) -> list[tuple[str, str]]:
    sources_raw = research.get("sources") or []
    out: list[tuple[str, str]] = []
    for s in sources_raw[:10]:
        if isinstance(s, dict):
            out.append((s.get("title") or s.get("url") or "", s.get("url") or ""))
        else:
            # Use a friendlier label by stripping protocol
            label = str(s).replace("https://", "").replace("http://", "").rstrip("/")
            out.append((label, str(s)))
    return [(l, u) for l, u in out if u]


# =====================================================================
# Public API: docx generation
# =====================================================================

def generate_docx(manifest: dict, research: dict, plan: dict, *,
                  slug: str, work_dir: str, portal: str) -> str:
    """Build the demo runbook .docx. Returns the path to the saved .docx."""
    doc = _build_doc(manifest, research, plan, slug=slug, work_dir=work_dir, portal=portal)
    out_path = os.path.join(work_dir, "demo-doc.docx")
    doc.save(out_path)
    return out_path


# =====================================================================
# Drive auth + upload + export
# =====================================================================

def _refresh_drive_access_token() -> str | None:
    """Refresh the gdrive OAuth access token via rclone's stored refresh token.

    Returns the access token string, or None if rclone is not configured /
    refresh fails. Never raises.
    """
    if not RCLONE_CONF_PATH.is_file():
        print(f"[doc_generator] WARN: {RCLONE_CONF_PATH} not found; skipping Drive upload",
              file=sys.stderr)
        return None
    cp = configparser.ConfigParser()
    try:
        cp.read(RCLONE_CONF_PATH)
    except Exception as e:
        print(f"[doc_generator] WARN: failed to parse rclone.conf: {e}", file=sys.stderr)
        return None
    if "gdrive" not in cp:
        print("[doc_generator] WARN: no [gdrive] section in rclone.conf", file=sys.stderr)
        return None
    try:
        tok = json.loads(cp["gdrive"]["token"])
    except Exception as e:
        print(f"[doc_generator] WARN: gdrive token unreadable: {e}", file=sys.stderr)
        return None
    refresh = tok.get("refresh_token")
    if not refresh:
        print("[doc_generator] WARN: no refresh_token in gdrive section", file=sys.stderr)
        return None
    try:
        req = urllib.request.Request(
            "https://oauth2.googleapis.com/token",
            data=urllib.parse.urlencode({
                "client_id": RCLONE_CLIENT_ID,
                "client_secret": RCLONE_CLIENT_SECRET,
                "refresh_token": refresh,
                "grant_type": "refresh_token",
            }).encode(),
            method="POST",
        )
        with urllib.request.urlopen(req, timeout=30) as r:
            return json.loads(r.read())["access_token"]
    except Exception as e:
        print(f"[doc_generator] WARN: token refresh failed: {e}", file=sys.stderr)
        return None


def upload_to_drive(docx_path: str, *, doc_title: str,
                    drive_folder_id: str = DEFAULT_DRIVE_FOLDER_ID,
                    replace_doc_id: str | None = None) -> dict:
    """Upload a .docx to Drive, converting to a Google Doc.

    If replace_doc_id is provided, PATCH that file (preserving its ID/URL).
    Otherwise POST a new file in the configured folder.

    PDF export is best-effort. On any error returns None values rather than raising.
    """
    result: dict = {"gdoc_url": None, "doc_id": None, "pdf_path": None}

    access_token = _refresh_drive_access_token()
    if not access_token:
        return result

    try:
        with open(docx_path, "rb") as f:
            docx_bytes = f.read()
    except Exception as e:
        print(f"[doc_generator] WARN: cannot read {docx_path}: {e}", file=sys.stderr)
        return result

    boundary = "------demoprepboundary"
    metadata: dict = {
        "mimeType": "application/vnd.google-apps.document",
        "name": doc_title,
    }
    if not replace_doc_id and drive_folder_id:
        metadata["parents"] = [drive_folder_id]

    body = (
        f"--{boundary}\r\n"
        "Content-Type: application/json; charset=UTF-8\r\n\r\n"
        f"{json.dumps(metadata)}\r\n"
        f"--{boundary}\r\n"
        "Content-Type: application/vnd.openxmlformats-officedocument.wordprocessingml.document\r\n\r\n"
    ).encode() + docx_bytes + f"\r\n--{boundary}--\r\n".encode()

    if replace_doc_id:
        url = (f"https://www.googleapis.com/upload/drive/v3/files/{replace_doc_id}"
               "?uploadType=multipart&supportsAllDrives=true")
        method = "PATCH"
    else:
        url = ("https://www.googleapis.com/upload/drive/v3/files"
               "?uploadType=multipart&supportsAllDrives=true")
        method = "POST"

    req = urllib.request.Request(
        url, data=body,
        headers={
            "Authorization": f"Bearer {access_token}",
            "Content-Type": f"multipart/related; boundary={boundary}",
        },
        method=method,
    )
    try:
        with urllib.request.urlopen(req, timeout=60) as r:
            resp = json.loads(r.read())
    except urllib.error.HTTPError as e:
        body_txt = ""
        try:
            body_txt = e.read().decode()
        except Exception:
            pass
        print(f"[doc_generator] WARN: Drive upload HTTP {e.code}: {body_txt[:300]}",
              file=sys.stderr)
        return result
    except Exception as e:
        print(f"[doc_generator] WARN: Drive upload failed: {e}", file=sys.stderr)
        return result

    doc_id = resp.get("id")
    if not doc_id:
        return result
    result["doc_id"] = doc_id
    result["gdoc_url"] = f"https://docs.google.com/document/d/{doc_id}/edit"

    # Best-effort PDF export
    pdf_out = os.path.join(os.path.dirname(docx_path), "demo-doc.pdf")
    pdf_path = export_pdf(doc_id, pdf_out)
    result["pdf_path"] = pdf_path
    return result


def export_pdf(doc_id: str, out_path: str) -> str | None:
    """Export the GDoc as PDF. Returns out_path on success, None on failure."""
    access_token = _refresh_drive_access_token()
    if not access_token:
        return None
    req = urllib.request.Request(
        f"https://www.googleapis.com/drive/v3/files/{doc_id}/export?mimeType=application/pdf",
        headers={"Authorization": f"Bearer {access_token}"},
    )
    try:
        with urllib.request.urlopen(req, timeout=60) as r, open(out_path, "wb") as f:
            f.write(r.read())
        return out_path
    except urllib.error.HTTPError as e:
        body_txt = ""
        try:
            body_txt = e.read().decode()
        except Exception:
            pass
        print(f"[doc_generator] WARN: PDF export HTTP {e.code}: {body_txt[:300]}",
              file=sys.stderr)
        return None
    except Exception as e:
        print(f"[doc_generator] WARN: PDF export failed: {e}", file=sys.stderr)
        return None


# =====================================================================
# CLI for local regen / smoke testing
# =====================================================================

def _main(argv: list[str]) -> int:
    if len(argv) < 2:
        print("usage: doc_generator.py <work_dir> [portal]", file=sys.stderr)
        return 2
    work_dir = argv[1]
    portal = argv[2] if len(argv) > 2 else "51393541"
    manifest = json.load(open(os.path.join(work_dir, "manifest.json")))
    research = json.load(open(os.path.join(work_dir, "research.json")))
    plan = json.load(open(os.path.join(work_dir, "build-plan.json")))
    slug = plan.get("slug") or os.path.basename(work_dir.rstrip("/")).replace("demo-prep-", "")
    out = generate_docx(manifest, research, plan, slug=slug, work_dir=work_dir, portal=portal)
    print(f"Saved {out} ({os.path.getsize(out)} bytes)")
    return 0


if __name__ == "__main__":
    sys.exit(_main(sys.argv))
